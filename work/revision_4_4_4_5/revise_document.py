# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BLACK = (20, 20, 20)
WHITE = (255, 255, 255)
LIGHT_BLUE = (232, 240, 248)
LIGHT_GRAY = (245, 245, 245)
BORDER_GRAY = (185, 185, 185)

FONT_SONG = Path("C:/Windows/Fonts/simsun.ttc")
FONT_HEI = Path("C:/Windows/Fonts/simhei.ttf")
FONT_YAHEI = Path("C:/Windows/Fonts/msyh.ttc")


class Diagram:
    def __init__(self, width: int, height: int, scale: int = 2):
        self.width = width
        self.height = height
        self.scale = scale
        self.image = Image.new("RGB", (width * scale, height * scale), WHITE)
        self.draw = ImageDraw.Draw(self.image)

    def _xy(self, values):
        return tuple(int(round(v * self.scale)) for v in values)

    def font(self, size: int, bold: bool = False):
        path = FONT_HEI if bold else FONT_SONG
        return ImageFont.truetype(str(path), size * self.scale)

    def line(self, points, width: int = 5, fill=BLACK):
        self.draw.line(
            [self._xy(point) for point in points],
            fill=fill,
            width=width * self.scale,
            joint="curve",
        )

    def arrow(self, points, width: int = 5, head: int = 18, fill=BLACK):
        self.line(points, width=width, fill=fill)
        p1 = points[-2]
        p2 = points[-1]
        dx = p2[0] - p1[0]
        dy = p2[1] - p1[1]
        length = math.hypot(dx, dy)
        if not length:
            return
        ux = dx / length
        uy = dy / length
        px = -uy
        py = ux
        base_x = p2[0] - ux * head
        base_y = p2[1] - uy * head
        half = head * 0.62
        triangle = [
            p2,
            (base_x + px * half, base_y + py * half),
            (base_x - px * half, base_y - py * half),
        ]
        self.draw.polygon([self._xy(point) for point in triangle], fill=fill)

    def dashed_arrow(self, start, end, width: int = 4, head: int = 16, dash: int = 16):
        x1, y1 = start
        x2, y2 = end
        dx = x2 - x1
        dy = y2 - y1
        length = math.hypot(dx, dy)
        ux = dx / length
        uy = dy / length
        position = 0.0
        gap = dash * 0.65
        while position < max(0, length - head):
            stop = min(position + dash, length - head)
            self.line(
                [
                    (x1 + ux * position, y1 + uy * position),
                    (x1 + ux * stop, y1 + uy * stop),
                ],
                width=width,
            )
            position = stop + gap
        self.arrow(
            [
                (x2 - ux * head, y2 - uy * head),
                (x2, y2),
            ],
            width=width,
            head=head,
        )

    def text(self, box, text: str, size: int = 40, bold: bool = False, fill=BLACK):
        font = self.font(size, bold=bold)
        scaled_box = self._xy(box)
        bbox = self.draw.multiline_textbbox(
            (0, 0), text, font=font, spacing=6 * self.scale, align="center"
        )
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        x = (scaled_box[0] + scaled_box[2] - text_width) / 2
        y = (scaled_box[1] + scaled_box[3] - text_height) / 2 - bbox[1]
        self.draw.multiline_text(
            (x, y),
            text,
            font=font,
            fill=fill,
            spacing=6 * self.scale,
            align="center",
        )

    def label(self, center, text: str, size: int = 34, bold: bool = False):
        font = self.font(size, bold=bold)
        x, y = self._xy(center)
        bbox = self.draw.textbbox((x, y), text, font=font, anchor="mm")
        pad = 5 * self.scale
        self.draw.rectangle(
            (bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad),
            fill=WHITE,
        )
        self.draw.text((x, y), text, font=font, fill=BLACK, anchor="mm")

    def rect(self, box, text: str, size: int = 40, radius: int = 0, fill=WHITE):
        scaled = self._xy(box)
        if radius:
            self.draw.rounded_rectangle(
                scaled,
                radius=radius * self.scale,
                fill=fill,
                outline=BLACK,
                width=5 * self.scale,
            )
        else:
            self.draw.rectangle(
                scaled,
                fill=fill,
                outline=BLACK,
                width=5 * self.scale,
            )
        self.text(box, text, size=size)

    def parallelogram(self, box, text: str, size: int = 40):
        x1, y1, x2, y2 = box
        skew = min(55, (x2 - x1) * 0.12)
        points = [(x1 + skew, y1), (x2, y1), (x2 - skew, y2), (x1, y2)]
        self.draw.polygon([self._xy(point) for point in points], fill=WHITE)
        self.draw.line(
            [self._xy(point) for point in points + [points[0]]],
            fill=BLACK,
            width=5 * self.scale,
            joint="curve",
        )
        self.text(box, text, size=size)

    def diamond(self, center, width: int, height: int, text: str, size: int = 38):
        x, y = center
        points = [
            (x, y - height / 2),
            (x + width / 2, y),
            (x, y + height / 2),
            (x - width / 2, y),
        ]
        self.draw.polygon([self._xy(point) for point in points], fill=WHITE)
        self.draw.line(
            [self._xy(point) for point in points + [points[0]]],
            fill=BLACK,
            width=5 * self.scale,
            joint="curve",
        )
        self.text((x - width / 2, y - height / 2, x + width / 2, y + height / 2), text, size=size)

    def border_and_title(self, title: str):
        self.draw.rectangle(
            self._xy((30, 70, self.width - 30, self.height - 45)),
            outline=BORDER_GRAY,
            width=3 * self.scale,
        )
        self.text((80, 85, self.width - 80, 180), title, size=52, bold=True)

    def save(self, path: Path, dpi=(300, 300)):
        path.parent.mkdir(parents=True, exist_ok=True)
        if self.scale != 1:
            self.image = self.image.resize(
                (self.width, self.height), Image.Resampling.LANCZOS
            )
        self.image.save(path, dpi=dpi)


def generate_account_flow(path: Path):
    d = Diagram(1500, 1980)
    d.border_and_title("账户管理流程图")
    d.rect((650, 210, 850, 290), "开始", radius=34)
    d.rect((600, 330, 900, 405), "进入账户页")
    d.diamond((750, 500), 320, 140, "选择操作")
    d.parallelogram((500, 620, 1000, 735), "填写账户信息或密码")
    d.rect((560, 790, 940, 885), "前端校验格式", radius=14)
    d.rect((560, 930, 940, 1025), "后端校验归属", radius=14)
    d.diamond((750, 1130), 380, 150, "校验通过")
    d.rect((170, 1065, 430, 1175), "提示错误\n返回修改", radius=14)
    d.rect((560, 1270, 940, 1365), "写入或查询数据", radius=14)
    d.rect((560, 1410, 940, 1505), "计算账户余额", radius=14)
    d.rect((560, 1550, 940, 1645), "刷新账户列表", radius=14)
    d.rect((650, 1700, 850, 1780), "结束", radius=34)

    d.arrow([(750, 290), (750, 330)])
    d.arrow([(750, 405), (750, 430)])
    d.arrow([(750, 570), (750, 620)])
    d.arrow([(750, 735), (750, 790)])
    d.arrow([(750, 885), (750, 930)])
    d.arrow([(750, 1025), (750, 1055)])
    d.arrow([(560, 1130), (430, 1130)])
    d.label((500, 1095), "否")
    d.arrow([(170, 1120), (100, 1120), (100, 678), (500, 678)])
    d.arrow([(750, 1205), (750, 1270)])
    d.label((785, 1235), "是")
    d.arrow([(750, 1365), (750, 1410)])
    d.arrow([(750, 1505), (750, 1550)])
    d.arrow([(750, 1645), (750, 1700)])
    d.save(path)


def generate_transaction_flow(path: Path):
    d = Diagram(1500, 1980)
    d.border_and_title("收支记录与转账流程图")
    d.rect((650, 210, 850, 290), "开始", radius=34)
    d.rect((600, 330, 900, 405), "进入收支页")
    d.diamond((750, 500), 360, 150, "选择业务")
    d.parallelogram((250, 650, 650, 770), "填写金额、账户\n分类、时间")
    d.parallelogram((850, 650, 1250, 770), "选择转入转出\n账户和金额")
    d.rect((300, 850, 600, 950), "校验分类和金额", radius=14)
    d.rect((900, 850, 1200, 950), "校验账户和金额", radius=14)
    d.diamond((450, 1080), 330, 140, "校验通过")
    d.diamond((1050, 1080), 330, 140, "校验通过")
    d.rect((70, 1025, 275, 1135), "提示错误", radius=14)
    d.rect((1225, 1025, 1430, 1135), "提示错误", radius=14)
    d.rect((300, 1220, 600, 1320), "写入收支流水", radius=14)
    d.rect((900, 1220, 1200, 1320), "生成转账流水", radius=14)
    d.rect((570, 1470, 930, 1570), "刷新列表和统计", radius=14)
    d.rect((650, 1640, 850, 1720), "结束", radius=34)

    d.arrow([(750, 290), (750, 330)])
    d.arrow([(750, 405), (750, 425)])
    d.arrow([(570, 500), (450, 500), (450, 650)])
    d.label((450, 610), "普通收支", size=32)
    d.arrow([(930, 500), (1050, 500), (1050, 650)])
    d.label((1050, 610), "账户转账", size=32)
    d.arrow([(450, 770), (450, 850)])
    d.arrow([(1050, 770), (1050, 850)])
    d.arrow([(450, 950), (450, 1010)])
    d.arrow([(1050, 950), (1050, 1010)])
    d.arrow([(285, 1080), (275, 1080)])
    d.label((300, 1040), "否", size=32)
    d.arrow([(70, 1080), (48, 1080), (48, 710), (250, 710)])
    d.arrow([(1215, 1080), (1225, 1080)])
    d.label((1200, 1040), "否", size=32)
    d.arrow([(1430, 1080), (1452, 1080), (1452, 710), (1250, 710)])
    d.arrow([(450, 1150), (450, 1220)])
    d.label((485, 1185), "是", size=32)
    d.arrow([(1050, 1150), (1050, 1220)])
    d.label((1085, 1185), "是", size=32)
    d.line([(450, 1320), (450, 1405), (750, 1405)])
    d.line([(1050, 1320), (1050, 1405), (750, 1405)])
    d.arrow([(750, 1405), (750, 1470)])
    d.arrow([(750, 1570), (750, 1640)])
    d.save(path)


def generate_budget_flow(path: Path):
    d = Diagram(1500, 1980)
    d.border_and_title("预算管理流程图")
    d.rect((650, 210, 850, 290), "开始", radius=34)
    d.rect((600, 330, 900, 405), "进入预算页")
    d.parallelogram((500, 455, 1000, 575), "选择月份、分类\n输入预算金额")
    d.rect((560, 630, 940, 730), "校验金额和归属", radius=14)
    d.diamond((750, 840), 400, 160, "同月分类\n已有预算")
    d.rect((260, 1020, 600, 1120), "更新预算记录", radius=14)
    d.rect((900, 1020, 1240, 1120), "新增预算记录", radius=14)
    d.rect((550, 1260, 950, 1360), "汇总分类支出金额", radius=14)
    d.rect((550, 1410, 950, 1510), "计算进度与预警", radius=14)
    d.rect((550, 1560, 950, 1660), "返回列表和进度", radius=14)
    d.rect((650, 1720, 850, 1800), "结束", radius=34)

    d.arrow([(750, 290), (750, 330)])
    d.arrow([(750, 405), (750, 455)])
    d.arrow([(750, 575), (750, 630)])
    d.arrow([(750, 730), (750, 760)])
    d.arrow([(550, 840), (430, 840), (430, 1020)])
    d.label((465, 805), "是", size=32)
    d.arrow([(950, 840), (1070, 840), (1070, 1020)])
    d.label((1035, 805), "否", size=32)
    d.line([(430, 1120), (430, 1195), (750, 1195)])
    d.line([(1070, 1120), (1070, 1195), (750, 1195)])
    d.arrow([(750, 1195), (750, 1260)])
    d.arrow([(750, 1360), (750, 1410)])
    d.arrow([(750, 1510), (750, 1560)])
    d.arrow([(750, 1660), (750, 1720)])
    d.save(path)


def generate_recurring_flow(path: Path):
    d = Diagram(1500, 1980)
    d.border_and_title("周期账单生成流程图")
    d.rect((650, 210, 850, 290), "开始", radius=34)
    d.rect((600, 330, 900, 405), "进入周期账单页")
    d.parallelogram((500, 455, 1000, 575), "维护名称、账户\n分类和周期")
    d.rect((560, 630, 940, 730), "校验金额和状态", radius=14)
    d.diamond((750, 850), 410, 170, "到期或手动\n生成")
    d.rect((250, 1040, 590, 1140), "保存模板等待", radius=14)
    d.rect((910, 1000, 1250, 1100), "生成交易记录", radius=14)
    d.rect((910, 1160, 1250, 1260), "更新下次日期", radius=14)
    d.rect((530, 1400, 970, 1500), "刷新流水和统计结果", radius=14)
    d.rect((650, 1570, 850, 1650), "结束", radius=34)

    d.arrow([(750, 290), (750, 330)])
    d.arrow([(750, 405), (750, 455)])
    d.arrow([(750, 575), (750, 630)])
    d.arrow([(750, 730), (750, 765)])
    d.arrow([(545, 850), (420, 850), (420, 1040)])
    d.label((455, 815), "否", size=32)
    d.arrow([(955, 850), (1080, 850), (1080, 1000)])
    d.label((1045, 815), "是", size=32)
    d.arrow([(1080, 1100), (1080, 1160)])
    d.line([(420, 1140), (420, 1335), (750, 1335)])
    d.line([(1080, 1260), (1080, 1335), (750, 1335)])
    d.arrow([(750, 1335), (750, 1400)])
    d.arrow([(750, 1500), (750, 1570)])
    d.save(path)


def generate_csv_flow(path: Path):
    d = Diagram(1500, 1980)
    d.border_and_title("CSV 导入流程图")
    d.rect((650, 210, 850, 290), "开始", radius=34)
    d.rect((600, 330, 900, 405), "进入导入页")
    d.parallelogram((500, 455, 1000, 575), "选择并上传 CSV 文件")
    d.diamond((750, 690), 400, 160, "文件格式\n正确")
    d.rect((240, 840, 540, 940), "提示格式错误", radius=14)
    d.rect((960, 840, 1260, 940), "解析 CSV 内容", radius=14)
    d.rect((900, 1000, 1320, 1110), "逐行校验账户、分类\n金额和日期", radius=14)
    d.diamond((1110, 1220), 380, 160, "数据逐行\n通过")
    d.rect((500, 1360, 800, 1460), "生成失败明细", radius=14)
    d.rect((500, 1510, 800, 1610), "导入有效数据", radius=14)
    d.rect((1150, 1435, 1450, 1535), "批量写入流水", radius=14)
    d.rect((560, 1680, 940, 1770), "生成导入结果", radius=14)
    d.rect((650, 1830, 850, 1910), "结束", radius=34)

    d.arrow([(750, 290), (750, 330)])
    d.arrow([(750, 405), (750, 455)])
    d.arrow([(750, 575), (750, 610)])
    d.arrow([(550, 690), (390, 690), (390, 840)])
    d.label((425, 650), "否", size=32)
    d.arrow([(950, 690), (1110, 690), (1110, 840)])
    d.label((1075, 650), "是", size=32)
    d.arrow([(1110, 940), (1110, 1000)])
    d.arrow([(1110, 1110), (1110, 1140)])
    d.arrow([(920, 1220), (840, 1220), (840, 1300), (650, 1300), (650, 1360)])
    d.label((860, 1180), "否", size=32)
    d.arrow([(1300, 1220), (1370, 1220), (1370, 1380), (1300, 1380), (1300, 1435)])
    d.label((1340, 1180), "是", size=32)
    d.arrow([(650, 1460), (650, 1510)])
    d.line([(390, 940), (160, 940), (160, 1640), (750, 1640)])
    d.line([(650, 1610), (650, 1640), (750, 1640)])
    d.line([(1300, 1535), (1300, 1640), (750, 1640)])
    d.arrow([(750, 1640), (750, 1680)])
    d.arrow([(750, 1770), (750, 1830)])
    d.save(path)


def generate_er_relationships(path: Path):
    d = Diagram(3000, 2344, scale=1)
    d.draw.rectangle((40, 110, 2960, 2260), outline=BORDER_GRAY, width=4)
    d.text((100, 15, 2900, 110), "数据库 E-R 关系图", size=72, bold=True)

    left_pairs = [
        ("user\n用户", "account\n账户", "拥有"),
        ("user\n用户", "transaction\n收支记录", "记录"),
        ("user\n用户", "budget\n预算", "设置"),
        ("user\n用户", "recurring_bill\n周期账单", "维护"),
        ("user\n用户", "budget_alert\n预算预警", "接收"),
        ("account\n账户", "transaction\n收支记录", "承载"),
    ]
    right_pairs = [
        ("account\n账户", "recurring_bill\n周期账单", "关联"),
        ("category\n分类", "transaction\n收支记录", "归类"),
        ("category\n分类", "budget\n预算", "限定"),
        ("category\n分类", "recurring_bill\n周期账单", "归类"),
        ("category\n分类", "budget_alert\n预算预警", "归类"),
        ("transaction\n转出流水", "transaction\n转入流水", "transfer_id 关联"),
    ]

    def pair(x1, x2, y, source, target, relation, ratio=("1", "N")):
        source_box = (x1, y - 65, x1 + 390, y + 65)
        target_box = (x2, y - 65, x2 + 390, y + 65)
        d.rect(source_box, source, size=42, fill=LIGHT_GRAY)
        d.rect(target_box, target, size=42, fill=LIGHT_GRAY)
        start = source_box[2]
        end = target_box[0]
        d.line([(start, y), (end, y)], width=5)
        d.label((start + 45, y - 28), ratio[0], size=38, bold=True)
        d.label((end - 45, y - 28), ratio[1], size=38, bold=True)
        d.label(((start + end) / 2, y - 40), relation, size=38)

    for index, item in enumerate(left_pairs):
        pair(100, 980, 300 + index * 300, *item)
    for index, item in enumerate(right_pairs):
        ratio = ("1", "1") if index == 5 else ("1", "N")
        pair(1580, 2460, 300 + index * 300, *item, ratio=ratio)

    d.text(
        (120, 2130, 2880, 2235),
        "说明：同名实体表示同一数据库对象；按关系分组绘制，以避免多条关联线交叉。",
        size=40,
    )
    d.save(path, dpi=(300, 300))


def generate_class_diagram(path: Path):
    d = Diagram(1800, 1350, scale=1)
    d.draw.rectangle((25, 85, 1775, 1325), outline=BORDER_GRAY, width=4)
    d.text((60, 5, 1740, 85), "核心业务类图", size=52, bold=True)

    modules = [
        {
            "name": "用户认证",
            "controller": ("UserController", ["+ register()", "+ login()"]),
            "service": ("UserServiceImpl", ["+ login()", "+ changePassword()"]),
            "mapper": ("UserMapper", ["+ selectByUsername()"]),
            "entity": ("User", ["- id, username", "- password, role"]),
        },
        {
            "name": "账户管理",
            "controller": ("AccountController", ["+ list()/create()", "+ getBalance()"]),
            "service": ("AccountServiceImpl", ["+ update()/delete()", "+ getBalance()"]),
            "mapper": ("AccountMapper", ["+ selectList()"]),
            "entity": ("Account", ["- userId, name", "- type, balance", "- currency, status"]),
        },
        {
            "name": "收支管理",
            "controller": ("Transaction\nController", ["+ create()/update()", "+ transfer()/importCsv()"]),
            "service": ("Transaction\nServiceImpl", ["+ transfer()", "+ importCsv()"]),
            "mapper": ("Transaction\nMapper", ["+ selectPage()"]),
            "entity": ("Transaction", ["- accountId", "- categoryId, type", "- amount, transferId"]),
        },
        {
            "name": "预算管理",
            "controller": ("BudgetController", ["+ save()/delete()", "+ getProgress()"]),
            "service": ("BudgetServiceImpl", ["+ getProgress()", "+ getAlert()"]),
            "mapper": ("BudgetMapper", ["+ selectByMonth()"]),
            "entity": ("Budget", ["- userId, categoryId", "- month, amount"]),
        },
        {
            "name": "周期账单",
            "controller": ("RecurringBill\nController", ["+ create()/update()", "+ generate()"]),
            "service": ("RecurringBill\nServiceImpl", ["+ deactivate()", "+ generate()"]),
            "mapper": ("RecurringBill\nMapper", ["+ selectActive()"]),
            "entity": ("RecurringBill", ["- accountId", "- categoryId, amount", "- period, dueDate"]),
        },
    ]

    lane_width = 320
    gap = 30
    x0 = 40

    def class_box(box, title, members, header_size=28, body_size=26):
        x1, y1, x2, y2 = box
        header_height = 55 if "\n" not in title else 76
        d.draw.rectangle(box, fill=WHITE, outline=BLACK, width=4)
        d.draw.rectangle((x1, y1, x2, y1 + header_height), fill=LIGHT_BLUE, outline=BLACK, width=4)
        d.text((x1 + 8, y1 + 4, x2 - 8, y1 + header_height - 4), title, size=header_size, bold=True)
        d.text((x1 + 10, y1 + header_height + 6, x2 - 10, y2 - 6), "\n".join(members), size=body_size)

    for index, module in enumerate(modules):
        x1 = x0 + index * (lane_width + gap)
        x2 = x1 + lane_width
        d.rect((x1, 105, x2, 165), module["name"], size=29, fill=LIGHT_GRAY)
        controller_box = (x1, 200, x2, 390)
        service_box = (x1, 460, x2, 675)
        mapper_box = (x1, 745, x2, 905)
        entity_box = (x1, 975, x2, 1150)
        class_box(controller_box, *module["controller"])
        class_box(service_box, *module["service"])
        class_box(mapper_box, *module["mapper"])
        class_box(entity_box, *module["entity"])
        center = (x1 + x2) / 2
        d.arrow([(center, 390), (center, 460)], width=4, head=15)
        d.arrow([(center, 675), (center, 745)], width=4, head=15)
        d.arrow([(center, 905), (center, 975)], width=4, head=15)

    d.text(
        (55, 1170, 1745, 1305),
        "共享支撑：StatisticsServiceImpl 基于 TransactionMapper 完成统计；AdminServiceImpl 基于 UserMapper 完成用户管理。",
        size=22,
    )
    d.save(path, dpi=(300, 300))


def generate_login_sequence(path: Path):
    d = Diagram(3000, 1900, scale=1)
    d.text((80, 10, 2920, 105), "登录核心时序图", size=62, bold=True)
    centers = [170, 610, 1050, 1490, 1930, 2370, 2810]
    names = ["用户", "LoginPage", "API 请求层", "UserController", "UserService", "UserMapper", "MySQL"]
    top = 125
    box_width = 330
    box_height = 90
    life_end = 1780
    for x, name in zip(centers, names):
        d.rect((x - box_width / 2, top, x + box_width / 2, top + box_height), name, size=42, fill=LIGHT_GRAY)
        y = top + box_height
        while y < life_end:
            stop = min(y + 22, life_end)
            d.line([(x, y), (x, stop)], width=3, fill=(90, 90, 90))
            y = stop + 16

    messages = [
        (0, 1, 300, "输入用户名和密码", False),
        (1, 2, 410, "login(data)", False),
        (2, 3, 520, "POST /api/v1/user/login", False),
        (3, 4, 630, "login(username, password)", False),
        (4, 5, 740, "selectByUsername", False),
        (5, 6, 850, "查询 user 表", False),
        (6, 5, 960, "返回用户记录", True),
        (5, 4, 1070, "返回 User", True),
        (4, 4, 1180, "BCrypt 校验并生成 JWT", False),
        (4, 3, 1320, "返回 LoginResponse", True),
        (3, 2, 1430, "Result.success(data)", True),
        (2, 1, 1540, "返回登录结果", True),
        (1, 0, 1650, "保存 token 并进入首页", True),
    ]
    for source, target, y, label, dashed in messages:
        if source == target:
            x = centers[source]
            d.line([(x, y), (x + 180, y), (x + 180, y + 75), (x, y + 75)], width=4)
            d.arrow([(x + 30, y + 75), (x, y + 75)], width=4, head=16)
            d.label((x + 215, y + 38), label, size=40)
            continue
        start = (centers[source], y)
        end = (centers[target], y)
        if dashed:
            d.dashed_arrow(start, end, width=4, head=16)
        else:
            d.arrow([start, end], width=4, head=16)
        d.label(((start[0] + end[0]) / 2, y - 28), label, size=40)

    d.save(path, dpi=(300, 300))


def set_run_font(run, size=10.5, bold=None, east_asia="宋体", latin="Times New Roman"):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def format_body_paragraph(paragraph, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def format_heading(paragraph, level: int):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.keep_with_next = True
    size = 14 if level == 2 else 12
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=False, east_asia="黑体", latin="Times New Roman")


def set_cell_text(cell, text: str, bold=False, center=False, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(13)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def prevent_row_split(row):
    row_properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_properties.append(cant_split)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def database_tables():
    return [
        (
            "user",
            "用户表",
            "用户表用于保存登录账号、密码哈希和角色信息，username 建立唯一索引，表结构如表 4-1 所示。",
            [
                ("id", "用户主键", "BIGINT", "主码，非空，自增"),
                ("username", "用户名", "VARCHAR(20)", "非空，唯一"),
                ("password", "BCrypt 密码哈希", "VARCHAR(100)", "非空"),
                ("role", "角色，0 普通用户，1 管理员", "TINYINT", "非空，默认 0"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "account",
            "账户表",
            "账户表保存用户的现金、银行卡、支付宝和微信账户，user_id 建立普通索引，表结构如表 4-2 所示。",
            [
                ("id", "账户主键", "BIGINT", "主码，非空，自增"),
                ("user_id", "所属用户 ID", "BIGINT", "非空，索引"),
                ("name", "账户名称", "VARCHAR(20)", "非空"),
                ("type", "账户类型，1 至 4", "TINYINT", "非空"),
                ("initial_balance", "初始余额", "DECIMAL(12,2)", "非空，默认 0.00"),
                ("currency", "币种代码", "VARCHAR(3)", "非空，默认 CNY"),
                ("status", "账户状态", "TINYINT", "非空，默认 1"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "category",
            "收支分类表",
            "收支分类表保存系统预置的收入和支出分类，业务层按 type 限定分类用途，表结构如表 4-3 所示。",
            [
                ("id", "分类主键", "BIGINT", "主码，非空，自增"),
                ("name", "分类名称", "VARCHAR(10)", "非空"),
                ("type", "分类类型，1 支出，2 收入", "TINYINT", "非空"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "transaction",
            "收支记录表",
            "收支记录表保存普通收支和转账流水，并对用户、账户、时间和 transfer_id 建立索引，表结构如表 4-4 所示。",
            [
                ("id", "收支记录主键", "BIGINT", "主码，非空，自增"),
                ("user_id", "所属用户 ID", "BIGINT", "非空，索引"),
                ("account_id", "关联账户 ID", "BIGINT", "非空，索引"),
                ("category_id", "关联分类 ID", "BIGINT", "非空"),
                ("type", "交易类型，1 收入，2 支出", "TINYINT", "非空"),
                ("amount", "交易金额", "DECIMAL(12,2)", "非空，大于 0"),
                ("note", "交易备注", "VARCHAR(200)", "允许为空"),
                ("time", "交易时间", "DATETIME", "非空，联合索引"),
                ("transfer_id", "转账关联 UUID", "VARCHAR(36)", "允许为空，索引"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "budget",
            "预算表",
            "预算表保存用户按月份和支出分类设置的预算，user_id、category_id、month 构成联合唯一索引，表结构如表 4-5 所示。",
            [
                ("id", "预算主键", "BIGINT", "主码，非空，自增"),
                ("user_id", "所属用户 ID", "BIGINT", "非空，联合唯一索引"),
                ("category_id", "预算分类 ID", "BIGINT", "非空，联合唯一索引"),
                ("month", "预算月份，YYYY-MM", "VARCHAR(7)", "非空，联合唯一索引"),
                ("amount", "预算金额", "DECIMAL(12,2)", "非空，大于 0"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "recurring_bill",
            "周期性账单表",
            "周期性账单表保存固定周期发生的收入或支出模板，用户、账户和分类字段均建立索引，表结构如表 4-6 所示。",
            [
                ("id", "周期账单主键", "BIGINT", "主码，非空，自增"),
                ("user_id", "所属用户 ID", "BIGINT", "非空，索引"),
                ("account_id", "关联账户 ID", "BIGINT", "非空，索引"),
                ("category_id", "关联分类 ID", "BIGINT", "非空，索引"),
                ("name", "账单名称", "VARCHAR(30)", "非空"),
                ("amount", "账单金额", "DECIMAL(12,2)", "非空，大于 0"),
                ("type", "类型，1 收入，2 支出", "TINYINT", "非空"),
                ("period", "daily/weekly/monthly/yearly", "VARCHAR(10)", "非空"),
                ("next_due_date", "下次到期日", "DATE", "非空"),
                ("status", "状态，1 活跃，0 停用", "TINYINT", "非空，默认 1"),
                ("create_time", "创建时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
        (
            "budget_alert",
            "预算预警表",
            "预算预警表保存定时计算得到的预算消耗与预警级别，并按用户、月份和分类建立查询索引，表结构如表 4-7 所示。",
            [
                ("id", "预算预警主键", "BIGINT", "主码，非空，自增"),
                ("user_id", "所属用户 ID", "BIGINT", "非空，联合索引"),
                ("category_id", "关联分类 ID", "BIGINT", "非空，索引"),
                ("month", "预算月份", "VARCHAR(7)", "非空，联合索引"),
                ("alert_level", "预警级别", "VARCHAR(20)", "非空"),
                ("budget_amount", "预算金额", "DECIMAL(12,2)", "非空"),
                ("spent_amount", "已消耗金额", "DECIMAL(12,2)", "非空"),
                ("percentage", "预算消耗百分比", "DECIMAL(8,2)", "非空"),
                ("create_time", "预警生成时间", "DATETIME", "非空，默认当前时间"),
                ("update_time", "更新时间", "DATETIME", "非空，自动更新"),
            ],
        ),
    ]


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_between(start_paragraph, end_paragraph):
    start = start_paragraph._p
    end = end_paragraph._p
    node = start.getnext()
    parent = start.getparent()
    while node is not None and node is not end:
        following = node.getnext()
        parent.remove(node)
        node = following
    if node is None:
        raise ValueError("End paragraph is not after start paragraph")


def move_before(element, anchor_paragraph):
    anchor_paragraph._p.addprevious(element)


def add_heading_before(document, anchor, text, level):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    format_heading(paragraph, level)
    move_before(paragraph._p, anchor)
    return paragraph


def add_body_before(document, anchor, text, first_line=True):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    format_body_paragraph(paragraph, first_line=first_line)
    move_before(paragraph._p, anchor)
    return paragraph


def add_caption_before(document, anchor, text, keep_with_next=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(16)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    move_before(paragraph._p, anchor)
    return paragraph


def add_picture_before(document, anchor, path, width_inches):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    move_before(paragraph._p, anchor)
    return paragraph


def add_relation_before(document, anchor, name, primary_key, remainder):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Pt(21)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{name}（")
    set_run_font(run)
    key_run = paragraph.add_run(primary_key)
    set_run_font(key_run)
    key_run.underline = True
    rest_run = paragraph.add_run(f"，{remainder}）")
    set_run_font(rest_run)
    move_before(paragraph._p, anchor)
    return paragraph


def add_table_before(
    document,
    anchor,
    headers,
    rows,
    widths,
    font_size=9,
    center_columns=(0, 2, 3),
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, center=True, size=font_size)
        set_cell_shading(table.rows[0].cells[index], "D9E2F3")
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_data):
            set_cell_text(
                row.cells[index],
                str(value),
                center=index in center_columns,
                size=font_size,
            )
    set_table_geometry(table, widths)
    move_before(table._tbl, anchor)
    return table


def add_module_table(document, anchor):
    rows = [
        ("用户认证", "UserController", "UserServiceImpl", "UserMapper", "注册、登录、改密、JWT 与 BCrypt 校验"),
        ("账户管理", "AccountController", "AccountServiceImpl", "AccountMapper", "账户增改停用、列表与实时余额计算"),
        ("分类查询", "Category\nController", "Category\nServiceImpl", "CategoryMapper", "读取系统预置的收入和支出分类"),
        ("收支管理", "Transaction\nController", "Transaction\nServiceImpl", "Transaction\nMapper", "流水增删改查、筛选、转账与 CSV 导入"),
        ("预算管理", "BudgetController", "BudgetServiceImpl", "BudgetMapper", "预算保存、进度计算与超支预警"),
        ("周期账单", "RecurringBill\nController", "RecurringBill\nServiceImpl", "RecurringBill\nMapper", "账单模板维护、停用与一键生成流水"),
        ("统计分析", "Statistics\nController", "Statistics\nServiceImpl", "Transaction\nMapper", "月度、年度、分类与趋势聚合"),
        ("管理员", "AdminController", "AdminServiceImpl", "UserMapper", "用户列表、删除用户与角色切换"),
    ]
    return add_table_before(
        document,
        anchor,
        ("模块", "控制层", "业务层", "数据访问层", "主要职责"),
        rows,
        (1100, 1700, 1900, 1600, 2300),
        font_size=8.5,
        center_columns=(0, 1, 2, 3),
    )


def add_page_break_before(document, anchor):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    move_before(paragraph._p, anchor)
    return paragraph


def replace_image_part(document, relation_id, image_path):
    part = document.part.related_parts[relation_id]
    part._blob = Path(image_path).read_bytes()


def build_document(input_path: Path, output_path: Path, work_dir: Path):
    flow_paths = {
        "rId11": work_dir / "flow_account.png",
        "rId12": work_dir / "flow_transaction.png",
        "rId13": work_dir / "flow_budget.png",
        "rId14": work_dir / "flow_recurring.png",
        "rId15": work_dir / "flow_csv.png",
    }
    er_path = work_dir / "er_relationships.png"
    class_path = work_dir / "class_diagram.png"
    sequence_path = work_dir / "login_sequence.png"

    generate_account_flow(flow_paths["rId11"])
    generate_transaction_flow(flow_paths["rId12"])
    generate_budget_flow(flow_paths["rId13"])
    generate_recurring_flow(flow_paths["rId14"])
    generate_csv_flow(flow_paths["rId15"])
    generate_er_relationships(er_path)
    generate_class_diagram(class_path)
    generate_login_sequence(sequence_path)

    document = Document(str(input_path))
    for relation_id, image_path in flow_paths.items():
        replace_image_part(document, relation_id, image_path)
    replace_image_part(document, "rId16", er_path)

    heading_44 = find_paragraph(document, "4.4 系统数据库设计")
    heading_46 = find_paragraph(document, "4.6 本章小结")
    remove_between(heading_44, heading_46)

    add_heading_before(document, heading_46, "4.4.1 数据库概念结构设计", 3)
    add_body_before(
        document,
        heading_46,
        "数据库概念结构设计以用户的财务活动为主线，围绕用户、账户、收支分类、收支记录、预算、周期账单和预算预警七个核心实体建立关系。用户实体保存账号与角色，是所有个人财务数据的归属主体；账户实体保存账户类型、初始余额、币种和状态，为流水记录提供资金载体。",
    )
    add_body_before(
        document,
        heading_46,
        "收支分类实体保存系统预置的收入、支出分类；收支记录实体保存账户、分类、金额、时间和备注等数据，是预算计算与统计分析的基础。账户转账在 transaction 表中生成一条支出流水和一条收入流水，两条记录使用相同的 transfer_id 建立逻辑关联。",
    )
    add_body_before(
        document,
        heading_46,
        "预算实体按照用户、分类和月份保存预算金额；周期账单实体保存固定周期发生的收入或支出模板；预算预警实体保存预算金额、实际支出、消耗百分比和预警级别。图 4-8 采用关系分组方式展示实体间的一对多关系，同名实体表示同一数据库对象，从而在完整表达关系的同时避免连线交叉。",
    )
    add_picture_before(document, heading_46, er_path, 6.0)
    add_caption_before(document, heading_46, "图 4-8 数据库 E-R 关系图")

    add_heading_before(document, heading_46, "4.4.2 数据库逻辑结构设计", 3)
    add_body_before(
        document,
        heading_46,
        "根据概念模型和实际业务约束，将核心实体转换为关系模式。各关系模式均满足字段原子性，重复的用户、账户和分类信息通过标识字段关联，避免在流水、预算和账单中重复存储基础资料。下列关系模式中带下划线的字段为主键。",
    )
    add_relation_before(document, heading_46, "用户 user", "id", "username，password，role，create_time，update_time")
    add_relation_before(document, heading_46, "账户 account", "id", "user_id，name，type，initial_balance，currency，status，create_time，update_time")
    add_relation_before(document, heading_46, "分类 category", "id", "name，type，create_time，update_time")
    add_relation_before(document, heading_46, "收支记录 transaction", "id", "user_id，account_id，category_id，type，amount，note，time，transfer_id，create_time，update_time")
    add_relation_before(document, heading_46, "预算 budget", "id", "user_id，category_id，month，amount，create_time，update_time")
    add_relation_before(document, heading_46, "周期账单 recurring_bill", "id", "user_id，account_id，category_id，name，amount，type，period，next_due_date，status，create_time，update_time")
    add_relation_before(document, heading_46, "预算预警 budget_alert", "id", "user_id，category_id，month，alert_level，budget_amount，spent_amount，percentage，create_time，update_time")
    add_body_before(
        document,
        heading_46,
        "系统未在 MySQL 中声明物理外键，而是在 Service 层根据当前登录用户校验 user_id、account_id 和 category_id 的归属与状态。这样既保持了逻辑关系完整，也避免级联删除误删个人财务流水。budget 表通过 user_id、category_id、month 联合唯一索引保证同一用户同月同分类仅有一条预算。",
    )

    add_heading_before(document, heading_46, "4.4.3 数据库物理结构设计", 3)
    add_body_before(
        document,
        heading_46,
        "数据库物理结构依据项目 sql/01-init.sql 与后端实体类整理。数据库采用 MySQL 8.4 LTS 和 InnoDB 引擎，字符集为 utf8mb4；金额字段使用 DECIMAL，Java 端使用 BigDecimal，避免浮点误差；日期时间字段分别使用 DATE 和 DATETIME。各核心表的字段、数据类型和完整性约束如下。",
    )
    for index, (table_name, chinese_name, intro, rows) in enumerate(database_tables(), start=1):
        intro_paragraph = add_body_before(document, heading_46, intro)
        intro_paragraph.paragraph_format.space_before = Pt(6)
        add_caption_before(
            document,
            heading_46,
            f"表 4-{index} {table_name} {chinese_name}",
            keep_with_next=True,
        )
        if index == 7:
            add_table_before(
                document,
                heading_46,
                ("字段名", "字段描述", "数据类型", "完整性约束"),
                rows[:6],
                (1450, 3000, 1650, 2500),
            )
            add_page_break_before(document, heading_46)
            add_caption_before(
                document,
                heading_46,
                "表 4-7 budget_alert 预算预警表（续）",
                keep_with_next=True,
            )
            add_table_before(
                document,
                heading_46,
                ("字段名", "字段描述", "数据类型", "完整性约束"),
                rows[6:],
                (1450, 3000, 1650, 2500),
            )
        else:
            add_table_before(
                document,
                heading_46,
                ("字段名", "字段描述", "数据类型", "完整性约束"),
                rows,
                (1450, 3000, 1650, 2500),
            )

    add_heading_before(document, heading_46, "4.5 详细模块设计", 2)
    add_heading_before(document, heading_46, "4.5.1 核心业务类设计", 3)
    add_body_before(
        document,
        heading_46,
        "系统后端采用 Controller、Service、Mapper、Entity 分层结构。Controller 接收 HTTP 请求并完成参数校验，Service 组织业务规则和事务，Mapper 通过 MyBatis-Plus 访问数据库，Entity 与数据表字段对应。核心业务类之间采用依赖调用关系，不把数据库操作直接写入控制层。",
    )
    add_body_before(
        document,
        heading_46,
        "图 4-9 展示用户认证、账户管理、收支管理、预算管理和周期账单五个核心模块的类协作关系。各模块由同名 Controller 调用 ServiceImpl，再由 ServiceImpl 调用 Mapper 并读写实体对象。StatisticsServiceImpl 复用 TransactionMapper 完成月度、年度、分类和趋势聚合；AdminServiceImpl 复用 UserMapper 完成用户管理。",
    )
    add_picture_before(document, heading_46, class_path, 6.0)
    add_caption_before(document, heading_46, "图 4-9 核心业务类图")

    add_heading_before(document, heading_46, "4.5.2 登录模块时序设计", 3)
    add_body_before(
        document,
        heading_46,
        "登录模块体现前后端分离系统的完整调用链。用户在 LoginPage 输入用户名和密码，页面通过 api/user.js 向 /api/v1/user/login 发送请求；UserController 接收请求后调用 UserServiceImpl，业务层通过 UserMapper 查询 user 表，并使用 BCrypt 校验密码。校验成功后生成 JWT 和 LoginResponse，再由统一 Result.success 返回前端。",
    )
    add_picture_before(document, heading_46, sequence_path, 6.0)
    add_caption_before(document, heading_46, "图 4-10 登录核心时序图")
    add_body_before(
        document,
        heading_46,
        "前端收到成功响应后将 token、用户 ID、用户名和角色写入 Pinia 状态及本地存储，并跳转到系统首页；后续 Axios 请求在请求头携带 token，后端 JWT 拦截器解析用户身份。用户名不存在或密码错误时，Service 抛出业务异常，统一异常处理器返回错误信息，前端保留在登录页并提示用户。",
    )

    add_heading_before(document, heading_46, "4.5.3 核心模块职责设计", 3)
    add_body_before(
        document,
        heading_46,
        "除登录模块外，各业务模块沿用相同分层方式。控制层只负责请求入口和响应包装，用户数据隔离、金额校验、账户状态检查、预算唯一性、转账事务和周期日期推进均由业务层完成。各模块对应的核心类与职责如表 4-8 所示。",
    )
    add_caption_before(document, heading_46, "表 4-8 核心模块类职责", keep_with_next=True)
    add_module_table(document, heading_46)
    add_body_before(
        document,
        heading_46,
        "收支转账由 TransactionServiceImpl 的事务方法一次生成转出和转入两条流水，任一写入失败时整体回滚；预算进度根据 budget.amount 与 transaction 支出聚合结果计算；周期账单生成流水后同步推进 next_due_date；CSV 导入逐行校验并返回成功数、失败数及失败原因。上述设计使控制层保持简洁，也便于单元测试直接验证业务层规则。",
    )

    summary = heading_46._p.getnext()
    if summary is not None and summary.tag == qn("w:p"):
        summary_paragraph = next(p for p in document.paragraphs if p._p is summary)
        summary_paragraph.clear()
        summary_paragraph.add_run(
            "本章完成了系统架构、总体功能模块、核心流程、数据库和详细模块设计。系统通过前后端分离降低耦合，通过五个核心流程图说明账户、收支、预算、周期账单和 CSV 导入过程；数据库设计以 E-R 关系、关系模式和七张物理表说明数据组织；详细设计通过核心业务类图、登录时序图和模块职责表明确 Controller、Service、Mapper 与 Entity 的协作方式。"
        )
        format_body_paragraph(summary_paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    check = Document(str(output_path))
    texts = "\n".join(paragraph.text for paragraph in check.paragraphs)
    required = [
        "4.4.1 数据库概念结构设计",
        "4.4.2 数据库逻辑结构设计",
        "4.4.3 数据库物理结构设计",
        "4.5.1 核心业务类设计",
        "4.5.2 登录模块时序设计",
        "4.5.3 核心模块职责设计",
        "图 4-10 登录核心时序图",
        "表 4-8 核心模块类职责",
    ]
    missing = [item for item in required if item not in texts]
    if missing:
        raise RuntimeError(f"Missing required content: {missing}")
    if len(check.tables) != 9:
        raise RuntimeError(f"Expected 9 tables, found {len(check.tables)}")
    if len(check.inline_shapes) != 13:
        raise RuntimeError(f"Expected 13 inline shapes, found {len(check.inline_shapes)}")
    return {
        "paragraphs": len(check.paragraphs),
        "tables": len(check.tables),
        "inline_shapes": len(check.inline_shapes),
        "output": str(output_path),
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--work-dir", type=Path, required=True)
    args = parser.parse_args()
    result = build_document(args.input, args.output, args.work_dir)
    for key, value in result.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
