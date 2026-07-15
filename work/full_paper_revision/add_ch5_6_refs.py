import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

path = Path(r'C:\Users\Administrator\Desktop\Question-12---Personal-Financial-Accounting-and-Analysis-System-master\选题标定-第12题-个人财务记账与分析系统\论文\个人财务记账与分析系统论文_2.2-7.2格式复验优化版.docx')
doc = Document(str(path))

# Map paragraph index -> new full text for body refs aligned with reference style
# Reference: "...页面如图5-1所示。" / "...如表6-1所示。"

replacements = {
    # 5.1 last body before fig 5-1 (para 192)
    192: None,  # will patch suffix
    197: None,
    202: None,
    209: None,
    214: None,
    219: None,
    225: None,
    226: None,
    244: None,
}

def set_para_text(p, text):
    # keep first run style if any
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)

# Build exact replacements from current text
def ensure_suffix(text, suffix):
    t = text.rstrip('。').rstrip()
    # remove existing similar suffix
    t = re.sub(r'(，)?如图\d+-\d+所示$', '', t)
    t = re.sub(r'(，)?如表\d+-\d+所示$', '', t)
    if not t.endswith('。') and not t.endswith('；'):
        return t + '，' + suffix + '。'
    # if ends with 。 already stripped; add
    return t + '，' + suffix + '。'

changes = []

# helper to get current full text
def full(p):
    return ''.join(r.text for r in p.runs) if p.runs else (p.text or '')

# Chapter 5 patches: last body paragraph before each figure/group
patches = {
    192: '如图5-1所示',
    197: '如图5-2所示',
    202: '收支记录列表页面如图5-3所示，账户转账页面如图5-4所示',
    209: '如图5-5所示',
    214: '如图5-6所示',
    219: '仪表盘统计页面如图5-7所示，财务分析页面如图5-8所示',
    225: '如图5-9所示',
    # para 226 is long covering account/recurring/import - need to read and add 5-10/5-11/5-12
}

# Read 226 fully
p226 = full(doc.paragraphs[226])
print('P226:', p226)

# Apply simple patches
for idx, tag in patches.items():
    old = full(doc.paragraphs[idx])
    if '如图5-' in old and tag.split('如图')[-1][:3] in old:
        # already has some
        if tag in old or tag.replace('如图','') in old:
            # ensure compact form
            new = re.sub(r'如图\s*(\d+-\d+)\s*所示', r'如图\1所示', old)
            if new != old:
                set_para_text(doc.paragraphs[idx], new)
                changes.append((idx, 'normalize', new[-40:]))
            continue
    new = ensure_suffix(old, tag if not tag.startswith('收支') and not tag.startswith('仪表') else tag)
    # for compound tags that already include 如图
    if tag.startswith('收支') or tag.startswith('仪表'):
        t = old.rstrip('。').rstrip()
        t = re.sub(r'(，)?如图\d+-\d+所示(，如图\d+-\d+所示)?$', '', t)
        new = t + '，' + tag + '。'
    set_para_text(doc.paragraphs[idx], new)
    changes.append((idx, old[-60:], new[-80:]))

# Special handling for 5.7 data extension section: paragraphs after 226 may be 227 etc
# dump 224-234
print('--- 5.7 area ---')
for i in range(224, 235):
    t = full(doc.paragraphs[i]).strip()
    if t:
        print(i, t[:160])

# We'll re-open after inspecting 5.7 structure from earlier dump:
# 225 settings, 226 extension, then images 5-9..5-12 at 228,230,232,234
# Need body before each image or one sentence covering all.

# Fix 226 to include 5-10/5-11/5-12 if it describes those modules
old226 = full(doc.paragraphs[226])
if '如图5-10' not in old226:
    t = old226.rstrip('。').rstrip()
    t = re.sub(r'(，)?如图5-\d+所示.*$', '', t)
    new226 = t + '。账户管理页面如图5-10所示，周期账单页面如图5-11所示，CSV导入页面如图5-12所示。'
    # ensure 5-9 is only on 225
    set_para_text(doc.paragraphs[226], new226)
    changes.append((226, old226[-50:], new226[-90:]))

# Ensure 225 ends with 如图5-9所示
old225 = full(doc.paragraphs[225])
if '如图5-9所示' not in old225:
    set_para_text(doc.paragraphs[225], ensure_suffix(old225, '如图5-9所示'))
    changes.append((225, old225[-50:], full(doc.paragraphs[225])[-60:]))
else:
    set_para_text(doc.paragraphs[225], re.sub(r'如图\s*(\d+-\d+)\s*所示', r'如图\1所示', old225))

# Chapter 6: para 244 introduce tables
old244 = full(doc.paragraphs[244])
if '如表6-1' not in old244:
    new244 = old244.rstrip('。').rstrip() + '。用户认证、账户管理、收支转账、预算管理、周期账单与CSV导入、统计分析与管理员测试用例分别如表6-1至表6-6所示。'
    set_para_text(doc.paragraphs[244], new244)
    changes.append((244, old244, new244))
else:
    set_para_text(doc.paragraphs[244], re.sub(r'如表\s*(\d+-\d+)\s*所示', r'如表\1所示', old244))

# Also mention in 6.4 briefly
old252 = full(doc.paragraphs[252])
if '表6-' not in old252:
    new252 = old252.rstrip('。').rstrip() + '。上述结论与表6-1至表6-6测试用例结果一致。'
    set_para_text(doc.paragraphs[252], new252)
    changes.append((252, old252[-40:], new252[-60:]))

doc.save(str(path))
print('CHANGES', len(changes))
for c in changes:
    print(c)

# verify refs
body = '\n'.join((p.text or '') for p in doc.paragraphs)
fig_refs = set(re.findall(r'如图(\d+-\d+)所示', body))
tab_refs = set(re.findall(r'如表(\d+-\d+)所示|表(\d+-\d+)至表(\d+-\d+)', body))
# expand range refs
range_refs = re.findall(r'表(\d+)-(\d+)至表(\d+)-(\d+)', body)
expanded=set()
for a,b,c,d in range_refs:
    if a==c:
        for n in range(int(b), int(d)+1):
            expanded.add(f'{a}-{n}')
fig_caps = set(re.findall(r'^图(\d+-\d+)\s', '\n'.join((p.text or '').strip() for p in doc.paragraphs), re.M))
# better
fig_caps=set(); tab_caps=set()
for p in doc.paragraphs:
    t=(p.text or '').strip()
    m=re.match(r'^图(\d+-\d+)\s', t)
    if m: fig_caps.add(m.group(1))
    m=re.match(r'^表(\d+-\d+)\s', t)
    if m: tab_caps.add(m.group(1))
print('fig refs', sorted(fig_refs))
print('missing fig refs', sorted(fig_caps-fig_refs))
# tables via range
print('tab expanded', sorted(expanded))
print('missing tab refs', sorted(tab_caps-expanded-set(re.findall(r'如表(\d+-\d+)所示', body))))
