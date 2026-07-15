import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = Path(r'C:\Users\Administrator\Desktop\Question-12---Personal-Financial-Accounting-and-Analysis-System-master\选题标定-第12题-个人财务记账与分析系统\论文\个人财务记账与分析系统论文_2.2-7.2格式复验优化版.docx')
doc = Document(str(path))
body = '\n'.join(p.text or '' for p in doc.paragraphs)

# all table captions
tabs=[]
figs=[]
for p in doc.paragraphs:
    t=(p.text or '').strip()
    m=re.match(r'^表(\d+-\d+)\s+(.+)$', t)
    if m: tabs.append(m.group(1)+' '+m.group(2))
    m=re.match(r'^图(\d+-\d+)\s+(.+)$', t)
    if m: figs.append(m.group(1)+' '+m.group(2))

# references: 如表x-x所示 / 表x-x / 如表x-x至表y-y
tab_direct = set(re.findall(r'如表(\d+-\d+)所示', body))
ranges = re.findall(r'表(\d+)-(\d+)至表(\d+)-(\d+)', body)
exp=set(tab_direct)
for a,b,c,d in ranges:
    if a==c:
        for n in range(int(b), int(d)+1):
            exp.add(f'{a}-{n}')
# also 核心环境如表3-1 形式
more = set(re.findall(r'表(\d+-\d+)', body))
# captions themselves match too; remove pure caption lines
cap_ids=set()
for p in doc.paragraphs:
    t=(p.text or '').strip()
    m=re.match(r'^表(\d+-\d+)\s', t)
    if m: cap_ids.add(m.group(1))

print('table captions', sorted(cap_ids))
print('body table id mentions (all)', sorted(more))
print('missing any caption id in body non-caption?', sorted(cap_ids - more))

# figure captions format check
bad=[]
for p in doc.paragraphs:
    t=(p.text or '').strip()
    if t.startswith('图') or t.startswith('表'):
        if re.match(r'^[图表]\s+\d', t):
            bad.append(('space', t))
        if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            bad.append(('align', t, p.alignment))
print('format bad', bad)

# spaced body
spaced=[]
for p in doc.paragraphs:
    t=p.text or ''
    if re.search(r'如图\s+\d|如表\s+\d', t):
        spaced.append(t[:100])
print('spaced body', spaced)

# image before each figure
miss=[]
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if not re.match(r'^图\d+-\d+', t):
        continue
    ok=False
    for j in (i-1,i-2):
        if j>=0 and (('w:drawing' in doc.paragraphs[j]._element.xml) or ('w:pict' in doc.paragraphs[j]._element.xml)):
            ok=True
    if not ok: miss.append(t)
print('missing images', miss)
print('figs', len(figs), 'tabs', len(tabs), 'tables objects', len(doc.tables))
print('file size', path.stat().st_size)
