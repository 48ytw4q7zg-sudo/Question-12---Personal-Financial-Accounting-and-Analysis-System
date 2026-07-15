# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import importlib.util
import math
import re
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(
    r"C:\Users\Administrator\Desktop\Question-12---Personal-Financial-Accounting-and-Analysis-System-master"
)
LEGACY_SCRIPT = ROOT / "work" / "revision_4_4_4_5" / "revise_document.py"
SPEC = importlib.util.spec_from_file_location("legacy_revision", LEGACY_SCRIPT)
if SPEC is None or SPEC.loader is None:
    raise RuntimeError(f"Cannot load helper module: {LEGACY_SCRIPT}")
legacy = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(legacy)


BLACK = (24, 24, 24)
WHITE = (255, 255, 255)
LIGHT_GRAY = (247, 247, 247)
FONT_SONG = Path("C:/Windows/Fonts/simsun.ttc")
FONT_HEI = Path("C:/Windows/Fonts/simhei.ttf")


class Diagram:
    def __init__(self, width: int, height: int, scale: int = 2):
        self.width = width
        self.height = height
        self.scale = scale
        self.image = Image.new("RGB", (width * scale, height * scale), WHITE)
        self.draw = ImageDraw.Draw(self.image)

    def xy(self, values):
        return tuple(int(round(value * self.scale)) for value in values)

    def font(self, size: int, bold: bool = False):
        font_path = FONT_HEI if bold else FONT_SONG
        return ImageFont.truetype(str(font_path), size * self.scale)

    def line(self, points, width: int = 3, fill=BLACK):
        self.draw.line(
            [self.xy(point) for point in points],
            fill=fill,
            width=width * self.scale,
            joint="curve",
        )

    def arrow(self, points, width: int = 3, head: int = 14, fill=BLACK):
        self.line(points, width=width, fill=fill)
        p1, p2 = points[-2], points[-1]
        dx, dy = p2[0] - p1[0], p2[1] - p1[1]
        length = math.hypot(dx, dy)
        if not length:
            return
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        bx, by = p2[0] - ux * head, p2[1] - uy * head
        half = head * 0.55
        triangle = [
            p2,
            (bx + px * half, by + py * half),
            (bx - px * half, by - py * half),
        ]
        self.draw.polygon([self.xy(point) for point in triangle], fill=fill)

    def dashed_arrow(self, points, width: int = 3, head: int = 14, dash: int = 12, hollow: bool = False):
        for start, end in zip(points, points[1:]):
            x1, y1 = start
            x2, y2 = end
            dx, dy = x2 - x1, y2 - y1
            length = math.hypot(dx, dy)
            if not length:
                continue
            ux, uy = dx / length, dy / length
            position = 0.0
            while position < length:
                stop = min(position + dash, length)
                self.line(
                    [
                        (x1 + ux * position, y1 + uy * position),
                        (x1 + ux * stop, y1 + uy * stop),
                    ],
                    width=width,
                )
                position = stop + dash * 0.7
        p1, p2 = points[-2], points[-1]
        dx, dy = p2[0] - p1[0], p2[1] - p1[1]
        length = math.hypot(dx, dy)
        if not length:
            return
        ux, uy = dx / length, dy / length
        if hollow:
            px, py = -uy, ux
            bx, by = p2[0] - ux * head, p2[1] - uy * head
            half = head * 0.55
            triangle = [
                p2,
                (bx + px * half, by + py * half),
                (bx - px * half, by - py * half),
            ]
            self.draw.polygon(
                [self.xy(point) for point in triangle],
                outline=BLACK,
                fill=WHITE,
                width=2 * self.scale,
            )
            self.line([(bx, by), p2], width=width)
        else:
            self.arrow(
                [(p2[0] - ux * head, p2[1] - uy * head), p2],
                width=width,
                head=head,
            )

    def text(self, box, text: str, size: int = 30, bold: bool = False, fill=BLACK):
        font = self.font(size, bold)
        scaled_box = self.xy(box)
        bbox = self.draw.multiline_textbbox(
            (0, 0), text, font=font, spacing=4 * self.scale, align="center"
        )
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        x = (scaled_box[0] + scaled_box[2] - text_width) / 2
        y = (scaled_box[1] + scaled_box[3] - text_height) / 2 - bbox[1]
        self.draw.multiline_text(
            (x, y),
            text,
            font=font,
            fill=fill,
            spacing=4 * self.scale,
            align="center",
        )

    def label(self, center, text: str, size: int = 24, bold: bool = False):
        font = self.font(size, bold)
        x, y = self.xy(center)
        bbox = self.draw.textbbox((x, y), text, font=font, anchor="mm")
        padding = 4 * self.scale
        self.draw.rectangle(
            (bbox[0] - padding, bbox[1] - padding, bbox[2] + padding, bbox[3] + padding),
            fill=WHITE,
        )
        self.draw.text((x, y), text, font=font, fill=BLACK, anchor="mm")

    def rect(self, box, text: str, size: int = 30, fill=WHITE, radius: int = 0):
        scaled = self.xy(box)
        if radius:
            self.draw.rounded_rectangle(
                scaled,
                radius=radius * self.scale,
                fill=fill,
                outline=BLACK,
                width=3 * self.scale,
            )
        else:
            self.draw.rectangle(scaled, fill=fill, outline=BLACK, width=3 * self.scale)
        self.text(box, text, size=size)

    def ellipse(self, box, text: str, size: int = 28, underline: bool = False):
        self.draw.ellipse(self.xy(box), fill=WHITE, outline=BLACK, width=3 * self.scale)
        self.text(box, text, size=size)
        if underline:
            font = self.font(size)
            bbox = self.draw.textbbox((0, 0), text, font=font)
            width = bbox[2] - bbox[0]
            x1 = ((box[0] + box[2]) * self.scale - width) / 2
            x2 = x1 + width
            y = ((box[1] + box[3]) / 2 + size * 0.55) * self.scale
            self.draw.line((x1, y, x2, y), fill=BLACK, width=2 * self.scale)

    def terminal(self, box, text: str, size: int = 30):
        radius = int((box[3] - box[1]) / 2)
        self.rect(box, text, size=size, radius=radius)

    def parallelogram(self, box, text: str, size: int = 28):
        x1, y1, x2, y2 = box
        skew = min(45, (x2 - x1) * 0.12)
        points = [(x1 + skew, y1), (x2, y1), (x2 - skew, y2), (x1, y2)]
        self.draw.polygon([self.xy(point) for point in points], fill=WHITE)
        self.draw.line(
            [self.xy(point) for point in points + [points[0]]],
            fill=BLACK,
            width=3 * self.scale,
            joint="curve",
        )
        self.text(box, text, size=size)

    def diamond(self, center, width: int, height: int, text: str, size: int = 27):
        x, y = center
        points = [
            (x, y - height / 2),
            (x + width / 2, y),
            (x, y + height / 2),
            (x - width / 2, y),
        ]
        self.draw.polygon([self.xy(point) for point in points], fill=WHITE)
        self.draw.line(
            [self.xy(point) for point in points + [points[0]]],
            fill=BLACK,
            width=3 * self.scale,
            joint="curve",
        )
        self.text(
            (x - width / 2, y - height / 2, x + width / 2, y + height / 2),
            text,
            size=size,
        )

    def data_store(self, box, text: str, code: str, size: int = 24):
        x1, y1, x2, y2 = box
        self.line([(x1, y1), (x2, y1)], width=3)
        self.line([(x1, y2), (x2, y2)], width=3)
        self.line([(x1 + 54, y1), (x1 + 54, y2)], width=3)
        self.text((x1, y1, x1 + 54, y2), code, size=size, bold=True)
        self.text((x1 + 54, y1, x2, y2), text, size=size)

    def actor(self, center, label: str, size: int = 24):
        x, y = center
        self.draw.ellipse(
            self.xy((x - 24, y - 95, x + 24, y - 47)),
            fill=WHITE,
            outline=BLACK,
            width=3 * self.scale,
        )
        self.line([(x, y - 47), (x, y + 45)], width=3)
        self.line([(x - 52, y - 8), (x + 52, y - 8)], width=3)
        self.line([(x, y + 45), (x - 45, y + 105)], width=3)
        self.line([(x, y + 45), (x + 45, y + 105)], width=3)
        self.text((x - 90, y + 112, x + 90, y + 165), label, size=size)

    def save(self, path: Path, dpi=(300, 300)):
        path.parent.mkdir(parents=True, exist_ok=True)
        if self.scale != 1:
            image = self.image.resize((self.width, self.height), Image.Resampling.LANCZOS)
        else:
            image = self.image
        image.save(path, dpi=dpi)


def generate_top_dfd(path: Path):
    d = Diagram(1800, 900)
    d.rect((75, 300, 320, 620), "普通用户", size=34)
    d.ellipse((625, 235, 1175, 685), "个人财务记账\n与分析系统", size=42)
    d.rect((1480, 300, 1725, 620), "管理员", size=34)

    left_flows = [
        (335, "注册登录、业务操作请求", True),
        (410, "账户、流水、预算与账单数据", True),
        (500, "余额、列表、预警与统计结果", False),
        (575, "导入结果与操作反馈", False),
    ]
    for y, label, to_system in left_flows:
        if to_system:
            d.arrow([(320, y), (625, y)])
        else:
            d.arrow([(625, y), (320, y)])
        d.label((472, y - 24), label, size=21)

    right_flows = [
        (375, "用户列表请求", False),
        (470, "删除与角色切换请求", False),
        (555, "用户数据与处理结果", True),
    ]
    for y, label, from_system in right_flows:
        if from_system:
            d.arrow([(1175, y), (1480, y)])
        else:
            d.arrow([(1480, y), (1175, y)])
        d.label((1328, y - 24), label, size=21)
    d.save(path)


def generate_level_one_dfd(path: Path):
    d = Diagram(2200, 1500, scale=1)
    d.rect((35, 600, 250, 860), "普通用户", size=30)
    d.rect((1850, 35, 2150, 225), "管理员", size=30)

    process_boxes = [
        (430, 105, 780, 265, "1.0\n身份认证与用户管理"),
        (430, 410, 780, 570, "2.0\n账户与收支处理"),
        (430, 760, 780, 920, "3.0\n预算与周期计划"),
        (430, 1120, 780, 1280, "4.0\n统计分析与 CSV 导入"),
    ]
    for box in process_boxes:
        d.ellipse(box[:4], box[4], size=26)

    stores = [
        (1050, 140, 1500, 230, "用户表", "D1"),
        (1050, 350, 1500, 440, "账户表", "D2"),
        (1050, 460, 1500, 550, "分类表", "D3"),
        (1050, 570, 1500, 660, "收支记录表", "D4"),
        (1050, 720, 1500, 810, "预算表", "D5"),
        (1050, 830, 1500, 920, "周期账单表", "D6"),
        (1050, 940, 1500, 1030, "预算预警表", "D7"),
        (1050, 1110, 1500, 1200, "分类表（同 D3）", "D3"),
        (1050, 1220, 1500, 1310, "收支记录表（同 D4）", "D4"),
        (1050, 1330, 1500, 1420, "预算表（同 D5）", "D5"),
    ]
    for x1, y1, x2, y2, label, code in stores:
        d.data_store((x1, y1, x2, y2), label, code)

    d.line([(250, 710), (335, 710), (335, 185)], width=3)
    d.line([(335, 185), (335, 1200)], width=3)
    flow_rows = [(185, "认证数据"), (490, "账户流水数据"), (840, "预算账单数据"), (1200, "统计导入数据")]
    for y, label in flow_rows:
        d.arrow([(335, y - 18), (430, y - 18)])
        d.arrow([(430, y + 22), (335, y + 22)])
        d.label((375, y - 43), label, size=19)

    d.arrow([(1850, 115), (820, 115), (820, 150), (780, 150)])
    d.label((1325, 88), "管理员登录与用户维护请求", size=20)
    d.arrow([(780, 220), (820, 220), (820, 255), (1850, 255), (1850, 225)])
    d.label((1325, 282), "用户列表与处理结果", size=20)

    store_links = [
        ((780, 165), (1050, 165), "both"),
        ((780, 205), (1050, 205), "both"),
        ((780, 450), (1050, 395), "both"),
        ((780, 490), (1050, 505), "read"),
        ((780, 530), (1050, 615), "both"),
        ((780, 800), (1050, 765), "both"),
        ((780, 840), (1050, 875), "both"),
        ((780, 880), (1050, 985), "both"),
        ((780, 1160), (1050, 1155), "read"),
        ((780, 1200), (1050, 1265), "both"),
        ((780, 1240), (1050, 1375), "read"),
    ]
    for start, end, mode in store_links:
        if mode in ("write", "both"):
            d.arrow([start, end])
        if mode in ("read", "both"):
            d.arrow([end, start])
    d.save(path)


def _use_case(d: Diagram, box, label: str, size: int = 26):
    d.ellipse(box, label, size=size)


def generate_user_use_case(path: Path):
    d = Diagram(1800, 2700, scale=1)
    d.actor((150, 1350), "普通用户", size=28)
    main_cases = [
        (260, "用户认证"),
        (560, "账户管理"),
        (900, "收支记录管理"),
        (1240, "预算管理"),
        (1580, "周期账单管理"),
        (1920, "统计分析"),
        (2250, "CSV 数据导入"),
        (2540, "个人设置"),
    ]
    for y, label in main_cases:
        _use_case(d, (520, y - 60, 900, y + 60), label)
        d.line([(205, 1342), (520, y)])

    subcases = {
        260: [(180, "用户注册"), (320, "用户登录")],
        560: [(490, "维护账户"), (650, "折算账户余额")],
        900: [(800, "维护收支流水"), (960, "筛选收支流水"), (1120, "账户转账")],
        1240: [(1250, "设置分类月预算"), (1410, "展示预算进度与预警")],
        1580: [(1570, "维护账单模板"), (1730, "生成收支流水")],
        1920: [(1910, "汇总收支数据"), (2070, "展示分析图表")],
        2250: [(2260, "校验并导入流水")],
        2540: [(2540, "修改登录密码")],
    }
    for main_y, items in subcases.items():
        for sub_y, label in items:
            _use_case(d, (1200, sub_y - 50, 1710, sub_y + 50), label, size=23)
            d.dashed_arrow(
                [(900, main_y), (1060, main_y), (1060, sub_y), (1200, sub_y)],
                hollow=True,
            )
            d.label((1085, sub_y - 22), "<<include>>", size=17)
    d.save(path)


def generate_admin_use_case(path: Path):
    d = Diagram(1700, 1100, scale=1)
    d.actor((155, 530), "管理员", size=28)
    main_cases = [
        (250, "管理员登录"),
        (520, "查看用户列表"),
        (800, "用户维护"),
    ]
    for y, label in main_cases:
        _use_case(d, (500, y - 65, 900, y + 65), label, size=27)
        d.line([(210, 520), (500, y)])

    maintenance = [(700, "删除用户"), (900, "切换用户角色")]
    for y, label in maintenance:
        _use_case(d, (1120, y - 58, 1560, y + 58), label, size=25)
        d.dashed_arrow(
            [(900, 800), (1020, 800), (1020, y), (1120, y)],
            hollow=True,
        )
        d.label((1045, y - 24), "<<include>>", size=18)
    d.save(path)


def generate_architecture(path: Path):
    d = Diagram(2100, 1450, scale=1)
    layers = [
        (60, 80, 2040, 260, "表现层", ["个人财务记账与分析系统 Web 界面"]),
        (
            60,
            300,
            2040,
            650,
            "应用层",
            [
                "注册登录", "账户管理", "分类浏览", "收支记录", "预算管理",
                "周期账单", "账户转账", "统计分析", "CSV 导入", "个人设置", "用户管理",
            ],
        ),
        (
            60,
            690,
            2040,
            910,
            "服务层",
            ["用户服务", "账户服务", "分类服务", "流水服务", "预算服务", "账单服务", "统计服务", "汇率服务", "管理员服务"],
        ),
        (
            60,
            950,
            2040,
            1160,
            "数据层",
            ["事务控制", "MyBatis-Plus 映射", "MySQL 持久化", "汇率参考缓存"],
        ),
        (
            60,
            1200,
            2040,
            1400,
            "基础架构层",
            ["Vue 3", "Element Plus", "Spring Boot", "JWT、BCrypt", "MySQL 8.4"],
        ),
    ]
    for x1, y1, x2, y2, layer_name, items in layers:
        d.rect((x1, y1, x2, y2), "", size=1)
        d.rect((90, y1 + 48, 315, y1 + 132), layer_name, size=28, radius=38)
        columns = 5 if len(items) >= 5 else len(items)
        rows = math.ceil(len(items) / columns)
        start_x = 390
        available = x2 - start_x - 40
        gap = 24
        box_width = (available - gap * (columns - 1)) / columns
        box_height = 68
        total_height = rows * box_height + (rows - 1) * 26
        start_y = y1 + (y2 - y1 - total_height) / 2
        for index, item in enumerate(items):
            row, column = divmod(index, columns)
            bx1 = start_x + column * (box_width + gap)
            by1 = start_y + row * (box_height + 26)
            d.rect((bx1, by1, bx1 + box_width, by1 + box_height), item, size=22)
    d.save(path)


def generate_module_tree(path: Path):
    d = Diagram(2500, 980, scale=1)
    root = (900, 40, 1600, 120)
    d.rect(root, "个人财务记账与分析系统", size=30)
    modules = [
        ("用户认证", ["用户注册", "用户登录", "修改密码"]),
        ("账户管理", ["新增账户", "编辑账户", "停用账户"]),
        ("收支管理", ["新增流水", "查询流水", "账户转账"]),
        ("预算管理", ["设置预算", "计算进度", "展示预警"]),
        ("周期账单", ["维护模板", "生成流水", "推进到期日"]),
        ("统计分析", ["汇总月年", "分类汇总", "趋势图表"]),
        ("数据及用户管理", ["导入数据", "查看用户", "维护用户"]),
    ]
    module_width = 280
    gap = 50
    total = len(modules) * module_width + (len(modules) - 1) * gap
    start_x = (2500 - total) / 2
    bus_y = 180
    root_cx = (root[0] + root[2]) / 2
    d.line([(root_cx, root[3]), (root_cx, bus_y)], width=3)
    left_cx = start_x + module_width / 2
    right_cx = start_x + total - module_width / 2
    d.line([(left_cx, bus_y), (right_cx, bus_y)], width=3)

    for index, (name, children) in enumerate(modules):
        x1 = start_x + index * (module_width + gap)
        x2 = x1 + module_width
        cx = (x1 + x2) / 2
        d.line([(cx, bus_y), (cx, 230)], width=3)
        d.rect((x1, 230, x2, 310), name, size=24)

        child_w = 80
        child_gap = 12
        child_total = len(children) * child_w + (len(children) - 1) * child_gap
        child_start = cx - child_total / 2
        child_bus = 360
        d.line([(cx, 310), (cx, child_bus)], width=3)
        if len(children) > 1:
            d.line(
                [
                    (child_start + child_w / 2, child_bus),
                    (child_start + child_total - child_w / 2, child_bus),
                ],
                width=3,
            )
        for c_index, child in enumerate(children):
            cx1 = child_start + c_index * (child_w + child_gap)
            cx2 = cx1 + child_w
            ccx = (cx1 + cx2) / 2
            d.line([(ccx, child_bus), (ccx, 400)], width=3)
            d.rect((cx1, 400, cx2, 520), child, size=16)
    d.save(path)


def generate_account_flow(path: Path):
    d = Diagram(1300, 1650)
    d.terminal((540, 45, 760, 120), "开始")
    d.rect((480, 165, 820, 245), "进入账户管理页面")
    d.diamond((650, 350), 360, 135, "选择账户操作")
    d.parallelogram((410, 455, 890, 555), "输入账户信息")
    d.rect((490, 610, 810, 690), "前端格式校验")
    d.diamond((650, 795), 360, 135, "格式是否正确")
    d.rect((65, 750, 340, 840), "提示格式错误")
    d.rect((490, 910, 810, 990), "后端归属与状态校验")
    d.diamond((650, 1095), 390, 140, "权限与状态是否有效")
    d.rect((960, 1050, 1235, 1140), "返回业务错误")
    d.rect((470, 1215, 830, 1300), "执行账户维护")
    d.rect((430, 1350, 870, 1435), "计算余额与多币种折算")
    d.rect((500, 1480, 800, 1555), "刷新账户列表")
    d.terminal((540, 1590, 760, 1645), "结束", size=26)

    d.arrow([(650, 120), (650, 165)])
    d.arrow([(650, 245), (650, 282)])
    d.arrow([(650, 418), (650, 455)])
    d.arrow([(650, 555), (650, 610)])
    d.arrow([(650, 690), (650, 728)])
    d.arrow([(470, 795), (340, 795)])
    d.label((405, 760), "否")
    d.arrow([(65, 795), (35, 795), (35, 505), (410, 505)])
    d.arrow([(650, 863), (650, 910)])
    d.label((686, 885), "是")
    d.arrow([(650, 990), (650, 1025)])
    d.arrow([(845, 1095), (960, 1095)])
    d.label((900, 1060), "否")
    d.arrow([(1235, 1095), (1265, 1095), (1265, 505), (890, 505)])
    d.arrow([(650, 1165), (650, 1215)])
    d.label((686, 1188), "是")
    d.arrow([(650, 1300), (650, 1350)])
    d.arrow([(650, 1435), (650, 1480)])
    d.arrow([(650, 1555), (650, 1590)])
    d.save(path)


def generate_transaction_flow(path: Path):
    d = Diagram(1500, 1700)
    d.terminal((640, 40, 860, 115), "开始")
    d.rect((575, 160, 925, 240), "进入收支记录页面")
    d.diamond((750, 350), 380, 140, "选择业务类型")
    d.parallelogram((120, 495, 590, 610), "填写账户、分类、金额与时间")
    d.parallelogram((910, 495, 1380, 610), "选择转出账户、转入账户与金额")
    d.rect((190, 690, 520, 780), "校验账户、分类与金额")
    d.rect((980, 690, 1310, 780), "校验账户归属与余额")
    d.diamond((355, 900), 350, 140, "普通收支是否合法")
    d.diamond((1145, 900), 350, 140, "转账参数是否合法")
    d.rect((85, 1035, 340, 1125), "返回校验错误")
    d.rect((1160, 1035, 1415, 1125), "返回校验错误")
    d.rect((185, 1205, 525, 1295), "写入一条收支流水")
    d.rect((930, 1180, 1360, 1270), "事务写入转出与转入流水")
    d.rect((565, 1410, 935, 1500), "刷新流水列表与统计")
    d.terminal((640, 1580, 860, 1655), "结束")

    d.arrow([(750, 115), (750, 160)])
    d.arrow([(750, 240), (750, 280)])
    d.arrow([(560, 350), (355, 350), (355, 495)])
    d.label((355, 450), "普通收支", size=22)
    d.arrow([(940, 350), (1145, 350), (1145, 495)])
    d.label((1145, 450), "账户转账", size=22)
    d.arrow([(355, 610), (355, 690)])
    d.arrow([(1145, 610), (1145, 690)])
    d.arrow([(355, 780), (355, 830)])
    d.arrow([(1145, 780), (1145, 830)])
    d.arrow([(180, 900), (85, 900), (85, 1035)])
    d.label((130, 865), "否")
    d.arrow([(85, 1080), (35, 1080), (35, 550), (120, 550)])
    d.arrow([(1320, 900), (1415, 900), (1415, 1035)])
    d.label((1370, 865), "否")
    d.arrow([(1415, 1080), (1465, 1080), (1465, 550), (1380, 550)])
    d.arrow([(355, 970), (355, 1205)])
    d.label((392, 1085), "是")
    d.arrow([(1145, 970), (1145, 1180)])
    d.label((1105, 1080), "是")
    d.line([(355, 1295), (355, 1350), (750, 1350)])
    d.line([(1145, 1270), (1145, 1350), (750, 1350)])
    d.arrow([(750, 1350), (750, 1410)])
    d.arrow([(750, 1500), (750, 1580)])
    d.save(path)


def generate_budget_flow(path: Path):
    d = Diagram(1300, 1600)
    d.terminal((540, 40, 760, 115), "开始")
    d.rect((480, 155, 820, 235), "进入预算管理页面")
    d.parallelogram((380, 285, 920, 400), "选择月份、支出分类并输入预算金额")
    d.rect((485, 450, 815, 535), "校验金额、分类与用户")
    d.diamond((650, 645), 380, 140, "输入是否合法")
    d.rect((70, 600, 340, 690), "提示校验错误")
    d.diamond((650, 835), 390, 145, "同月同分类是否已有预算")
    d.rect((150, 970, 470, 1060), "更新原预算记录")
    d.rect((830, 970, 1150, 1060), "新增预算记录")
    d.rect((460, 1170, 840, 1260), "汇总该分类当月支出")
    d.rect((440, 1310, 860, 1400), "计算进度与四级预警")
    d.rect((490, 1450, 810, 1530), "返回预算列表与结果")
    d.terminal((540, 1550, 760, 1595), "结束", size=24)

    d.arrow([(650, 115), (650, 155)])
    d.arrow([(650, 235), (650, 285)])
    d.arrow([(650, 400), (650, 450)])
    d.arrow([(650, 535), (650, 575)])
    d.arrow([(460, 645), (340, 645)])
    d.label((400, 610), "否")
    d.arrow([(70, 645), (35, 645), (35, 343), (380, 343)])
    d.arrow([(650, 715), (650, 762)])
    d.label((686, 740), "是")
    d.arrow([(455, 835), (310, 835), (310, 970)])
    d.label((350, 800), "是")
    d.arrow([(845, 835), (990, 835), (990, 970)])
    d.label((950, 800), "否")
    d.line([(310, 1060), (310, 1110), (650, 1110)])
    d.line([(990, 1060), (990, 1110), (650, 1110)])
    d.arrow([(650, 1110), (650, 1170)])
    d.arrow([(650, 1260), (650, 1310)])
    d.arrow([(650, 1400), (650, 1450)])
    d.arrow([(650, 1530), (650, 1550)])
    d.save(path)


def generate_recurring_flow(path: Path):
    d = Diagram(1300, 1640)
    d.terminal((540, 40, 760, 115), "开始")
    d.rect((455, 155, 845, 235), "进入周期账单页面")
    d.parallelogram(
        (355, 285, 945, 410),
        "填写名称、账户、分类、金额、周期和到期日",
        size=24,
    )
    d.rect((470, 465, 830, 550), "校验账户、分类与日期")
    d.diamond((650, 665), 390, 140, "账单信息是否合法")
    d.rect((65, 620, 340, 710), "提示校验错误")
    d.rect((475, 785, 825, 870), "保存或更新账单模板")
    d.diamond((650, 990), 390, 145, "是否一键生成流水")
    d.rect((110, 1115, 430, 1205), "返回账单列表")
    d.rect((850, 1090, 1190, 1180), "事务写入收支流水")
    d.rect((850, 1240, 1190, 1330), "推进下次到期日至未来")
    d.rect((465, 1440, 835, 1525), "刷新账单、流水与统计")
    d.terminal((540, 1570, 760, 1635), "结束")

    d.arrow([(650, 115), (650, 155)])
    d.arrow([(650, 235), (650, 285)])
    d.arrow([(650, 410), (650, 465)])
    d.arrow([(650, 550), (650, 595)])
    d.arrow([(455, 665), (340, 665)])
    d.label((400, 630), "否")
    d.arrow([(65, 665), (35, 665), (35, 347), (355, 347)])
    d.arrow([(650, 735), (650, 785)])
    d.label((686, 760), "是")
    d.arrow([(650, 870), (650, 918)])
    d.arrow([(455, 990), (270, 990), (270, 1115)])
    d.label((315, 955), "否")
    d.arrow([(845, 990), (1020, 990), (1020, 1090)])
    d.label((980, 955), "是")
    d.arrow([(1020, 1180), (1020, 1240)])
    d.line([(270, 1205), (270, 1380), (650, 1380)])
    d.line([(1020, 1330), (1020, 1380), (650, 1380)])
    d.arrow([(650, 1380), (650, 1440)])
    d.arrow([(650, 1525), (650, 1570)])
    d.save(path)


def generate_csv_flow(path: Path):
    d = Diagram(1400, 1800)
    d.terminal((590, 35, 810, 110), "开始")
    d.rect((515, 150, 885, 230), "进入数据导入页面")
    d.parallelogram((385, 280, 1015, 400), "选择账户并上传 CSV 文件")
    d.diamond((700, 520), 430, 150, "文件大小与扩展名是否有效")
    d.rect((70, 475, 370, 565), "提示文件格式错误")
    d.rect((535, 655, 865, 740), "OpenCSV 解析文件")
    d.rect((535, 790, 865, 875), "读取下一行数据")
    d.diamond((700, 1000), 390, 145, "是否仍有数据行")
    d.diamond((700, 1160), 390, 145, "当前行是否合法")
    d.rect((180, 1280, 500, 1370), "记录失败行与原因")
    d.rect((900, 1280, 1240, 1370), "写入流水并累计成功数")
    d.rect((505, 1510, 895, 1600), "返回成功数、失败数和明细")
    d.terminal((590, 1690, 810, 1765), "结束")

    d.arrow([(700, 110), (700, 150)])
    d.arrow([(700, 230), (700, 280)])
    d.arrow([(700, 400), (700, 445)])
    d.arrow([(485, 520), (370, 520)])
    d.label((425, 485), "否")
    d.arrow([(70, 520), (35, 520), (35, 1727), (590, 1727)])
    d.arrow([(700, 595), (700, 655)])
    d.label((738, 625), "是")
    d.arrow([(700, 740), (700, 790)])
    d.arrow([(700, 875), (700, 928)])
    d.arrow([(505, 1000), (80, 1000), (80, 1555), (505, 1555)])
    d.label((390, 965), "否")
    d.arrow([(700, 1072), (700, 1088)])
    d.label((738, 1080), "是")
    d.arrow([(505, 1160), (340, 1160), (340, 1280)])
    d.label((420, 1125), "否")
    d.arrow([(895, 1160), (1070, 1160), (1070, 1280)])
    d.label((980, 1125), "是")
    d.line([(340, 1370), (340, 1430), (1340, 1430)])
    d.line([(1070, 1370), (1070, 1430), (1340, 1430)])
    d.arrow([(1340, 1430), (1340, 835), (865, 835)])
    d.arrow([(700, 1600), (700, 1690)])
    d.save(path)


def generate_entity_attribute(path: Path, entity: str, attributes: list[str], primary_key: str):
    count = len(attributes)
    width, height = 1900, 900
    d = Diagram(width, height, scale=1)
    center_x, center_y = width / 2, height / 2
    entity_box = (center_x - 150, center_y - 65, center_x + 150, center_y + 65)
    radius_x = 700
    radius_y = 330
    positions = []
    for index in range(count):
        angle = -math.pi / 2 + 2 * math.pi * index / count
        positions.append(
            (
                center_x + radius_x * math.cos(angle),
                center_y + radius_y * math.sin(angle),
            )
        )
    for x, y in positions:
        d.line([(center_x, center_y), (x, y)], width=3)
    d.rect(entity_box, entity, size=31)
    for attribute, (x, y) in zip(attributes, positions):
        d.ellipse(
            (x - 145, y - 48, x + 145, y + 48),
            attribute,
            size=22,
            underline=attribute == primary_key,
        )
    d.save(path)


def generate_er_diagram(path: Path):
    # Each Chen relationship has an independent connector; shared buses are avoided.
    d = Diagram(3200, 1900, scale=1)

    user = (1450, 40, 1750, 140)
    d.rect(user, "用户", size=34)

    mid_y1, mid_y2 = 760, 870
    boxes = {
        "tx": (120, mid_y1, 470, mid_y2),
        "acc": (640, mid_y1, 990, mid_y2),
        "rb": (1160, mid_y1, 1540, mid_y2),
        "bud": (1710, mid_y1, 2060, mid_y2),
        "alert": (2240, mid_y1, 2620, mid_y2),
        "cat": (1450, 1680, 1750, 1780),
    }
    labels = {
        "tx": "收支记录",
        "acc": "账户",
        "rb": "周期账单",
        "bud": "预算",
        "alert": "预算预警",
        "cat": "收支分类",
    }
    for key, box in boxes.items():
        d.rect(box, labels[key], size=31)

    def top_mid(box):
        return ((box[0] + box[2]) / 2, box[1])

    def bottom_mid(box):
        return ((box[0] + box[2]) / 2, box[3])

    def left_mid(box):
        return (box[0], (box[1] + box[3]) / 2)

    def right_mid(box):
        return (box[2], (box[1] + box[3]) / 2)

    # User -> middle entities. Distinct source points and routing lanes keep every
    # relationship visually independent.
    top_links = [
        ("tx", "记录", 295, 420, 1480, 180),
        ("acc", "拥有", 815, 420, 1560, 240),
        ("rb", "维护", 1350, 420, 1640, 300),
        ("bud", "设置", 1885, 420, 1720, 360),
    ]
    for key, label, x, dy, source_x, lane_y in top_links:
        target = top_mid(boxes[key])
        d.line([(source_x, user[3]), (source_x, lane_y), (x, lane_y), (x, dy - 55)])
        d.line([(x, dy + 55), target])
        d.diamond((x, dy), 200, 100, label, size=24)
        d.label((source_x + 22, user[3] + 24), "1", size=21)
        d.label((x + 24, (dy + target[1]) / 2), "N", size=21)

    # Horizontal middle relations
    horiz = [
        (right_mid(boxes["tx"]), left_mid(boxes["acc"]), "承载", "N", "1"),
        (right_mid(boxes["acc"]), left_mid(boxes["rb"]), "关联", "1", "N"),
        (right_mid(boxes["bud"]), left_mid(boxes["alert"]), "生成", "1", "N"),
    ]
    for start, end, label, sc, ec in horiz:
        mx = (start[0] + end[0]) / 2
        my = start[1]
        d.line([start, (mx - 100, my)])
        d.line([(mx + 100, my), end])
        d.diamond((mx, my), 200, 100, label, size=24)
        d.label((start[0] + 28, my - 34), sc, size=21)
        d.label((end[0] - 28, my - 34), ec, size=21)

    # Category bottom links. Each path has its own lane and its own entry point
    # on the category entity, so no relationship shares a connector segment.
    cat = boxes["cat"]
    bottom_links = [
        ("tx", "归类", 295, 1180, [(295, 1705), (cat[0], 1705)]),
        ("rb", "归类", 1350, 1180, [(1350, 1470), (1510, 1470), (1510, cat[1])]),
        ("bud", "限定", 1885, 1180, [(1885, 1430), (1690, 1430), (1690, cat[1])]),
        ("alert", "归类", 2430, 1180, [(2430, 1755), (cat[2], 1755)]),
    ]
    for key, label, x, dy, route in bottom_links:
        target = bottom_mid(boxes[key])
        d.line([target, (x, dy - 55)])
        d.diamond((x, dy), 200, 100, label, size=24)
        d.line([(x, dy + 55), *route])
        d.label((x + 24, (target[1] + dy) / 2), "N", size=21)
        d.label((x + 24, dy + 80), "1", size=21)

    d.save(path, dpi=(300, 300))


def _class_box(d: Diagram, box, title: str, attributes: list[str], methods: list[str]):
    x1, y1, x2, y2 = box
    title_height = 58
    attribute_height = max(110, len(attributes) * 30 + 20)
    d.draw.rectangle(box, fill=WHITE, outline=BLACK, width=3)
    d.draw.line((x1, y1 + title_height, x2, y1 + title_height), fill=BLACK, width=3)
    d.draw.line(
        (x1, y1 + title_height + attribute_height, x2, y1 + title_height + attribute_height),
        fill=BLACK,
        width=3,
    )
    d.text((x1, y1, x2, y1 + title_height), title, size=26, bold=True)
    d.text(
        (x1 + 8, y1 + title_height + 5, x2 - 8, y1 + title_height + attribute_height - 5),
        "\n".join(attributes),
        size=18,
    )
    d.text(
        (x1 + 8, y1 + title_height + attribute_height + 5, x2 - 8, y2 - 5),
        "\n".join(methods),
        size=18,
    )


def generate_class_diagram(path: Path):
    d = Diagram(3200, 2050, scale=1)
    boxes = {
        "User": (1360, 40, 1840, 430),
        "Transaction": (80, 700, 560, 1220),
        "Account": (680, 700, 1160, 1190),
        "RecurringBill": (1280, 700, 1820, 1280),
        "Budget": (2010, 700, 2490, 1190),
        "BudgetAlert": (2640, 700, 3160, 1220),
        "Category": (1360, 1600, 1840, 2000),
    }
    _class_box(d, boxes["User"], "User 用户", ["- id: Long", "- username: String", "- password: String", "- role: Integer"], ["+ register()", "+ login()", "+ changePassword()"])
    _class_box(d, boxes["Transaction"], "Transaction 收支记录", ["- id: Long", "- userId: Long", "- accountId: Long", "- categoryId: Long", "- type: Integer", "- amount: BigDecimal", "- transferId: String"], ["+ create()", "+ update()", "+ transfer()", "+ importCsv()"])
    _class_box(d, boxes["Account"], "Account 账户", ["- id: Long", "- userId: Long", "- name: String", "- type: Integer", "- initialBalance: BigDecimal", "- currency: String", "- status: Integer"], ["+ create()", "+ update()", "+ disable()", "+ getBalance()"])
    _class_box(d, boxes["RecurringBill"], "RecurringBill 周期账单", ["- id: Long", "- userId: Long", "- accountId: Long", "- categoryId: Long", "- amount: BigDecimal", "- period: String", "- nextDueDate: LocalDate", "- status: Integer"], ["+ create()", "+ update()", "+ deactivate()", "+ generate()"])
    _class_box(d, boxes["Budget"], "Budget 预算", ["- id: Long", "- userId: Long", "- categoryId: Long", "- month: String", "- amount: BigDecimal"], ["+ save()", "+ delete()", "+ getProgress()"])
    _class_box(d, boxes["BudgetAlert"], "BudgetAlert 预算预警", ["- id: Long", "- userId: Long", "- categoryId: Long", "- month: String", "- alertLevel: String", "- spentAmount: BigDecimal", "- percentage: BigDecimal"], ["+ calculateLevel()", "+ persist()", "+ listAlerts()"])
    _class_box(d, boxes["Category"], "Category 收支分类", ["- id: Long", "- name: String", "- type: Integer", "- createTime: LocalDateTime"], ["+ listByType()", "+ validateType()"])

    user_sources = [(1400, 430), (1500, 430), (1600, 430), (1700, 430), (1800, 430)]
    top_targets = [
        ((320, 700), "1", "N"),
        ((920, 700), "1", "N"),
        ((1550, 700), "1", "N"),
        ((2250, 700), "1", "N"),
        ((2900, 700), "1", "N"),
    ]
    for source, (target, source_card, target_card) in zip(user_sources, top_targets):
        d.line([source, target], width=3)
        d.label(((source[0] * 0.82 + target[0] * 0.18), (source[1] * 0.82 + target[1] * 0.18)), source_card, size=20)
        d.label(((source[0] * 0.18 + target[0] * 0.82), (source[1] * 0.18 + target[1] * 0.82)), target_card, size=20)

    category_sources = [(1400, 1600), (1500, 1600), (1700, 1600), (1800, 1600)]
    bottom_targets = [
        ((320, 1220), "1", "N"),
        ((1550, 1280), "1", "N"),
        ((2250, 1190), "1", "N"),
        ((2900, 1220), "1", "N"),
    ]
    for source, (target, source_card, target_card) in zip(category_sources, bottom_targets):
        d.line([source, target], width=3)
        d.label(((source[0] * 0.82 + target[0] * 0.18), (source[1] * 0.82 + target[1] * 0.18)), source_card, size=20)
        d.label(((source[0] * 0.18 + target[0] * 0.82), (source[1] * 0.18 + target[1] * 0.82)), target_card, size=20)

    d.line([(560, 960), (680, 960)], width=3)
    d.label((590, 925), "N", size=20)
    d.label((650, 925), "1", size=20)
    d.line([(1160, 960), (1280, 960)], width=3)
    d.label((1190, 925), "1", size=20)
    d.label((1250, 925), "N", size=20)
    d.line([(2490, 960), (2640, 960)], width=3)
    d.label((2520, 925), "1", size=20)
    d.label((2610, 925), "N", size=20)
    d.save(path, dpi=(300, 300))


def generate_sequence(path: Path, participants: list[str], messages: list[tuple[int, int, str, bool]]):
    width = max(2500, 360 * len(participants) + 120)
    height = 1750
    d = Diagram(width, height, scale=1)
    margin = 110
    span = width - margin * 2
    centers = [margin + index * span / (len(participants) - 1) for index in range(len(participants))]
    box_width = min(300, span / len(participants) * 0.82)
    top = 60
    box_height = 82
    life_end = 1670
    for participant_index, (x, participant) in enumerate(zip(centers, participants)):
        if participant_index == 0 and participant in {"用户", "普通用户", "管理员"}:
            d.actor((x, 125), participant, size=24)
            y = 305
        else:
            d.rect(
                (x - box_width / 2, top, x + box_width / 2, top + box_height),
                participant,
                size=24,
            )
            y = top + box_height
        while y < life_end:
            stop = min(y + 18, life_end)
            d.line([(x, y), (x, stop)], width=2, fill=(95, 95, 95))
            y = stop + 13

    start_y = 360
    step = min(105, (life_end - start_y - 40) / max(1, len(messages)))
    for index, (source, target, label, is_return) in enumerate(messages):
        y = start_y + index * step
        if source == target:
            x = centers[source]
            d.line([(x, y), (x + 115, y), (x + 115, y + 50), (x, y + 50)], width=3)
            d.arrow([(x + 28, y + 50), (x, y + 50)], width=3, head=12)
            d.label((x + 160, y + 25), label, size=21)
            continue
        start = (centers[source], y)
        end = (centers[target], y)
        if is_return:
            d.dashed_arrow([start, end], width=2, head=12, dash=10)
        else:
            d.arrow([start, end], width=3, head=12)
        d.label(((start[0] + end[0]) / 2, y - 20), label, size=21)
    d.save(path, dpi=(300, 300))


def generate_login_sequence(path: Path):
    generate_sequence(
        path,
        ["用户", "LoginPage", "Axios", "UserController", "UserService", "UserMapper", "MySQL"],
        [
            (0, 1, "输入用户名和密码", False),
            (1, 2, "login(data)", False),
            (2, 3, "POST /user/login", False),
            (3, 4, "login(username,password)", False),
            (4, 5, "selectByUsername", False),
            (5, 6, "查询 user 表", False),
            (6, 5, "返回用户记录", True),
            (5, 4, "返回 User", True),
            (4, 4, "BCrypt 校验并生成 JWT", False),
            (4, 3, "LoginResponse", True),
            (3, 2, "Result.success", True),
            (2, 1, "返回 token 与角色", True),
            (1, 0, "保存状态并进入首页", True),
        ],
    )


def generate_transfer_sequence(path: Path):
    generate_sequence(
        path,
        ["用户", "TransferPage", "Axios", "TransactionController", "TransactionService", "Mapper", "MySQL"],
        [
            (0, 1, "选择转出账户与转入账户并输入金额", False),
            (1, 2, "transfer(data)", False),
            (2, 3, "POST /transaction/transfer", False),
            (3, 4, "transfer(userId,request)", False),
            (4, 5, "校验账户归属与余额", False),
            (5, 6, "查询 account", False),
            (6, 5, "返回账户", True),
            (4, 4, "开启 @Transactional", False),
            (4, 5, "insert 转出与转入流水", False),
            (5, 6, "写入 transaction 两条记录", False),
            (6, 5, "写入成功", True),
            (5, 4, "返回流水", True),
            (4, 3, "TransferDTO", True),
            (3, 2, "Result.success", True),
            (2, 1, "刷新流水与余额", True),
        ],
    )


def generate_budget_sequence(path: Path):
    generate_sequence(
        path,
        ["用户", "BudgetPage", "Axios", "BudgetController", "BudgetService", "Mapper", "MySQL"],
        [
            (0, 1, "选择月份、分类与金额", False),
            (1, 2, "saveBudget(data)", False),
            (2, 3, "POST /budget", False),
            (3, 4, "save(userId,request)", False),
            (4, 5, "查询同月同分类预算", False),
            (5, 6, "SELECT budget", False),
            (6, 5, "返回预算记录", True),
            (4, 5, "insert 或 update budget", False),
            (5, 6, "持久化预算", False),
            (4, 5, "汇总 transaction 支出", False),
            (5, 6, "SUM(amount)", False),
            (6, 5, "返回已支出金额", True),
            (4, 3, "预算进度与预警级别", True),
            (3, 2, "Result.success", True),
            (2, 1, "展示进度条和预警", True),
        ],
    )


def generate_recurring_sequence(path: Path):
    generate_sequence(
        path,
        ["用户", "RecurringBillPage", "Axios", "RecurringBillController", "RecurringBillService", "Mapper", "MySQL"],
        [
            (0, 1, "点击一键生成", False),
            (1, 2, "generate(id)", False),
            (2, 3, "POST /recurring-bill/{id}/generate", False),
            (3, 4, "generate(userId,id)", False),
            (4, 5, "查询账单、账户与分类", False),
            (5, 6, "SELECT recurring/account/category", False),
            (6, 5, "返回关联数据", True),
            (4, 4, "开启 @Transactional", False),
            (4, 5, "insert transaction", False),
            (5, 6, "写入收支流水", False),
            (4, 4, "推进 nextDueDate 至未来", False),
            (4, 5, "update recurring_bill", False),
            (5, 6, "更新到期日", False),
            (4, 3, "TransactionDTO", True),
            (3, 2, "Result.success", True),
            (2, 1, "刷新账单和流水", True),
        ],
    )


def generate_csv_sequence(path: Path):
    generate_sequence(
        path,
        ["用户", "ImportPage", "Axios", "TransactionController", "TransactionService", "OpenCSV、Mapper", "MySQL"],
        [
            (0, 1, "选择账户并上传 CSV", False),
            (1, 2, "importCsv(file,accountId)", False),
            (2, 3, "POST /transaction/import", False),
            (3, 4, "importCsv(userId,file,accountId)", False),
            (4, 4, "校验 5MB、扩展名和账户归属", False),
            (4, 5, "CSVReader.readNext()", False),
            (5, 4, "返回一行字段", True),
            (4, 4, "校验列数、类型、金额和日期", False),
            (4, 5, "insert 有效流水", False),
            (5, 6, "写入 transaction", False),
            (6, 5, "写入成功", True),
            (4, 4, "累计成功数、失败数与失败原因", False),
            (4, 3, "ImportResultDTO", True),
            (3, 2, "Result.success", True),
            (2, 1, "展示导入结果", True),
        ],
    )


def set_table_borders(table, top_size=14, bottom_size=10):
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("left", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")
    for edge, size in (("top", top_size), ("bottom", bottom_size)):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_cell_bottom_border(cell, size=10):
    cell_properties = cell._tc.get_or_add_tcPr()
    borders = cell_properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell_properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def add_three_line_table_before(
    document,
    anchor,
    headers,
    rows,
    widths,
    font_size=8.5,
    center_columns=(),
    keep_together=False,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for index, header in enumerate(headers):
        legacy.set_cell_text(
            table.rows[0].cells[index], header, bold=False, center=True, size=font_size
        )
        set_cell_bottom_border(table.rows[0].cells[index])
    legacy.set_repeat_header(table.rows[0])
    legacy.prevent_row_split(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        legacy.prevent_row_split(row)
        for index, value in enumerate(row_data):
            legacy.set_cell_text(
                row.cells[index],
                str(value),
                center=index in center_columns,
                size=font_size,
            )
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    legacy.set_table_geometry(table, widths, indent=120)
    if keep_together:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    anchor._p.addprevious(table._tbl)
    return table


def compact_figure_table_refs(text: str) -> str:
    text = re.sub(r"([图表])\s*(\d+-\d+)", r"\1\2", text)
    text = re.sub(r"(图\d+-\d+)\s*所示", r"\1所示", text)
    text = re.sub(r"(表\d+-\d+)\s*所示", r"\1所示", text)
    text = re.sub(r"(表\d+-\d+)\s*至\s*(表\d+-\d+)", r"\1至\2", text)
    return text


def ensure_section_properties(document):
    if document.element.body.xpath("./w:sectPr"):
        return
    template = Document()
    section = template.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    document.element.body.append(deepcopy(section._sectPr))


def format_chapter_headings(document):
    chapter_titles = (
        "1 绪论",
        "2 可行性分析与需求分析",
        "3 系统相关理论和技术",
        "4 系统设计",
        "5 系统实现",
        "6 系统测试",
        "7 总结和展望",
    )
    for title in chapter_titles:
        paragraph = legacy.find_paragraph(document, title)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(18)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.page_break_before = True
        for run in paragraph.runs:
            legacy.set_run_font(
                run,
                size=16,
                bold=False,
                east_asia="黑体",
                latin="Times New Roman",
            )


def add_figure(document, anchor, path: Path, caption: str, width: float = 6.0):
    legacy.add_picture_before(document, anchor, path, width)
    legacy.add_caption_before(document, anchor, compact_figure_table_refs(caption))


def add_table_caption(document, anchor, caption: str):
    return legacy.add_caption_before(
        document, anchor, compact_figure_table_refs(caption), keep_with_next=True
    )


def remove_after(paragraph):
    node = paragraph._p.getnext()
    parent = paragraph._p.getparent()
    while node is not None:
        next_node = node.getnext()
        parent.remove(node)
        node = next_node


def generate_all_figures(output_dir: Path):
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "top_dfd": output_dir / "fig_2_1_top_dfd.png",
        "level_one_dfd": output_dir / "fig_2_2_level_one_dfd.png",
        "user_use_case": output_dir / "fig_2_3_user_use_case.png",
        "admin_use_case": output_dir / "fig_2_4_admin_use_case.png",
        "architecture": output_dir / "fig_4_1_architecture.png",
        "module_tree": output_dir / "fig_4_2_module_tree.png",
        "flow_account": output_dir / "fig_4_3_flow_account.png",
        "flow_transaction": output_dir / "fig_4_4_flow_transaction.png",
        "flow_budget": output_dir / "fig_4_5_flow_budget.png",
        "flow_recurring": output_dir / "fig_4_6_flow_recurring.png",
        "flow_csv": output_dir / "fig_4_7_flow_csv.png",
        "attr_user": output_dir / "fig_4_8_attr_user.png",
        "attr_account": output_dir / "fig_4_9_attr_account.png",
        "attr_category": output_dir / "fig_4_10_attr_category.png",
        "attr_transaction": output_dir / "fig_4_11_attr_transaction.png",
        "attr_budget": output_dir / "fig_4_12_attr_budget.png",
        "attr_recurring": output_dir / "fig_4_13_attr_recurring.png",
        "attr_alert": output_dir / "fig_4_14_attr_alert.png",
        "er": output_dir / "fig_4_15_er.png",
        "class": output_dir / "fig_4_16_class.png",
        "seq_login": output_dir / "fig_4_17_seq_login.png",
        "seq_transfer": output_dir / "fig_4_18_seq_transfer.png",
        "seq_budget": output_dir / "fig_4_19_seq_budget.png",
        "seq_recurring": output_dir / "fig_4_20_seq_recurring.png",
        "seq_csv": output_dir / "fig_4_21_seq_csv.png",
    }
    generate_top_dfd(paths["top_dfd"])
    generate_level_one_dfd(paths["level_one_dfd"])
    generate_user_use_case(paths["user_use_case"])
    generate_admin_use_case(paths["admin_use_case"])
    generate_architecture(paths["architecture"])
    generate_module_tree(paths["module_tree"])
    generate_account_flow(paths["flow_account"])
    generate_transaction_flow(paths["flow_transaction"])
    generate_budget_flow(paths["flow_budget"])
    generate_recurring_flow(paths["flow_recurring"])
    generate_csv_flow(paths["flow_csv"])
    generate_entity_attribute(
        paths["attr_user"],
        "用户",
        ["用户ID", "用户名", "密码", "角色", "创建时间", "更新时间"],
        "用户ID",
    )
    generate_entity_attribute(
        paths["attr_account"],
        "账户",
        ["账户ID", "用户ID", "账户名称", "账户类型", "初始余额", "币种", "状态", "创建时间", "更新时间"],
        "账户ID",
    )
    generate_entity_attribute(
        paths["attr_category"],
        "收支分类",
        ["分类ID", "分类名称", "分类类型", "创建时间", "更新时间"],
        "分类ID",
    )
    generate_entity_attribute(
        paths["attr_transaction"],
        "收支记录",
        ["记录ID", "用户ID", "账户ID", "分类ID", "交易类型", "金额", "备注", "交易时间", "转账关联ID", "创建时间", "更新时间"],
        "记录ID",
    )
    generate_entity_attribute(
        paths["attr_budget"],
        "预算",
        ["预算ID", "用户ID", "分类ID", "预算月份", "预算金额", "创建时间", "更新时间"],
        "预算ID",
    )
    generate_entity_attribute(
        paths["attr_recurring"],
        "周期账单",
        ["账单ID", "用户ID", "账户ID", "分类ID", "账单名称", "金额", "类型", "周期", "下次到期日", "状态", "创建时间", "更新时间"],
        "账单ID",
    )
    generate_entity_attribute(
        paths["attr_alert"],
        "预算预警",
        ["预警ID", "用户ID", "分类ID", "预算月份", "预警级别", "预算金额", "已支出金额", "消耗百分比", "创建时间", "更新时间"],
        "预警ID",
    )
    generate_er_diagram(paths["er"])
    generate_class_diagram(paths["class"])
    generate_login_sequence(paths["seq_login"])
    generate_transfer_sequence(paths["seq_transfer"])
    generate_budget_sequence(paths["seq_budget"])
    generate_recurring_sequence(paths["seq_recurring"])
    generate_csv_sequence(paths["seq_csv"])
    return paths


def add_body(document, anchor, text: str, first_line: bool = True):
    return legacy.add_body_before(
        document, anchor, compact_figure_table_refs(text), first_line=first_line
    )


def add_heading(document, anchor, text: str, level: int):
    return legacy.add_heading_before(document, anchor, text, level)


def add_page_break(document, anchor):
    return legacy.add_page_break_before(document, anchor)


def build_abstract_and_early_sections(document):
    heading_abstract = legacy.find_paragraph(document, "摘要")
    heading_1 = legacy.find_paragraph(document, "1 绪论")
    legacy.remove_between(heading_abstract, heading_1)

    add_body(
        document,
        heading_1,
        "随着移动支付、网络购物和多账户资金管理逐渐普及，个人收支数据分散在现金、银行卡及第三方支付账户中。仅依靠纸质记录或电子表格难以持续维护账户余额、分类支出、月度预算和周期账单，也不便于从历史流水中获得直观的财务分析结果。为解决个人财务数据分散、统计滞后和预算执行不透明等问题，本文设计并实现了一套个人财务记账与分析系统。",
    )
    add_body(
        document,
        heading_1,
        "系统采用 B/S 架构和前后端分离模式。前端使用 Vue 3、Vue Router、Pinia、Element Plus、Axios 和 ECharts 构建页面、管理登录状态并展示统计图表；后端使用 Spring Boot、MyBatis-Plus 和 MySQL，按照 Controller、Service、Mapper、Entity 分层实现 REST 接口、业务规则和数据持久化，并结合 JWT、BCrypt、OpenCSV、统一异常处理及事务机制保证身份认证、密码存储、文件导入和跨表操作的可靠性。",
    )
    add_body(
        document,
        heading_1,
        "系统实现了用户注册登录、账户管理、分类浏览、收支流水增删改查、账户转账、预算设置与分级预警、周期账单生成、月度与年度统计、多维图表、CSV 导入、密码修改以及管理员用户维护等功能。账户、流水、预算、周期账单和预警数据均按当前用户隔离；转账和周期账单生成使用事务保证数据一致性。实际运行页面、前端生产构建和后端测试结果表明，系统能够完成个人财务记账、汇总、预算控制与分析的主要业务流程。",
    )
    add_body(
        document,
        heading_1,
        "关键词：个人财务；收支记账；预算预警；Spring Boot；Vue 3",
        first_line=False,
    )

    heading_2 = legacy.find_paragraph(document, "2 可行性分析与需求分析")
    legacy.remove_between(heading_1, heading_2)
    add_body(
        document,
        heading_2,
        "个人财务管理系统既要解决收支数据记录问题，也要处理账户余额、预算进度、周期日期、统计口径和用户数据隔离等相互关联的业务规则。本章从项目背景、研究意义和主要研究内容三个方面说明系统建设的必要性与实现范围。",
    )
    add_heading(document, heading_2, "1.1 研究背景及意义", 2)
    add_heading(document, heading_2, "1.1.1 研究背景", 3)
    add_body(
        document,
        heading_2,
        "数字化支付提高了消费和资金流转效率，也使个人财务记录呈现账户多、交易频率高、数据来源分散等特点。同一用户可能同时使用现金、银行卡、支付宝、微信及不同币种账户，交易信息分布在多个平台，难以形成统一、连续且可追溯的个人财务视图。",
    )
    add_body(
        document,
        heading_2,
        "传统手工记账和电子表格能够保存基础流水，但账户余额通常需要人工回算，分类统计、月度预算和周期账单也依赖重复操作。当发生账户转账、历史账单补生成或批量导入时，若缺少统一的数据模型、事务控制和校验规则，容易出现重复记录、余额口径不一致、预算进度滞后及错误数据混入等问题。",
    )
    add_body(
        document,
        heading_2,
        "因此，有必要构建面向个人用户的财务记账与分析系统，将账户、分类、流水、预算、周期账单和统计分析集中到同一平台，并通过身份认证、数据隔离和标准化业务流程提高记录效率与数据可信度。",
    )
    add_heading(document, heading_2, "1.1.2 研究意义", 3)
    add_body(
        document,
        heading_2,
        "从使用价值看，系统可以统一保存不同账户和币种的收支信息，实时计算账户当前余额，通过预算进度和分级预警提示消费风险，并利用分类占比、月度趋势和年度对比帮助用户了解资金流向，为日常支出控制和财务决策提供数据依据。",
    )
    add_body(
        document,
        heading_2,
        "从工程实践看，个人财务数据对金额精度、访问权限和跨表一致性具有明确要求。系统将 DECIMAL 与 BigDecimal 金额计算、JWT 身份认证、BCrypt 密码哈希、用户维度数据隔离、事务回滚、统一异常处理和逐行文件校验应用到具体业务中，可验证前后端分离架构在个人财务场景中的实现方法。",
    )
    add_body(
        document,
        heading_2,
        "从教学与扩展价值看，系统覆盖需求分析、架构设计、数据库设计、业务实现、界面实现和软件测试的完整开发过程，形成了较清晰的模块边界和数据关系，可为后续扩展实时汇率、账单适配、消费预测和资产负债分析提供基础。",
    )
    add_heading(document, heading_2, "1.2 主要研究内容", 2)
    add_body(
        document,
        heading_2,
        "本文主要研究如何依据个人财务业务需求完成系统设计与实现。系统面向普通用户和管理员两类角色：普通用户完成注册登录、账户维护、分类浏览、收支记录、账户转账、预算与预警、周期账单、统计分析、CSV 导入和密码修改；管理员在独立权限校验后查看用户列表、删除用户并切换用户角色。",
    )
    add_body(
        document,
        heading_2,
        "在前端实现方面，采用 Vue 3 组件化开发页面，使用 Vue Router 组织 12 个路由并实施路由守卫，使用 Pinia 保存令牌和用户角色，使用 Axios 统一封装请求与异常响应，使用 Element Plus 完成表单、表格、对话框、上传和进度展示，使用 ECharts 呈现分类金额、收支趋势和年度对比。",
    )
    add_body(
        document,
        heading_2,
        "在后端与数据实现方面，采用 Spring Boot 和 MyBatis-Plus 建立 Controller、Service、Mapper、Entity 分层，使用 MySQL 保存 user、account、category、transaction、budget、recurring_bill 和 budget_alert 七类核心数据。系统通过 JWT 与 BCrypt 实现认证，通过 user_id 条件实现数据隔离，通过事务处理账户转账和周期账单生成，通过 OpenCSV 完成批量流水导入。",
    )
    add_body(
        document,
        heading_2,
        "在分析、设计与验证方面，本文使用数据流图和用例图确定需求边界，使用架构图、功能模块图和流程图描述处理过程，使用实体属性图、E-R 图、类图和时序图说明数据及对象协作关系，并结合实际运行页面、功能测试、异常测试、构建测试和后端自动化测试核验实现结果。",
    )

    heading_22 = legacy.find_paragraph(document, "2.2 系统需求分析")
    legacy.remove_between(heading_2, heading_22)
    add_body(
        document,
        heading_22,
        "系统开发前需要从经济投入、用户操作和技术实现三个方面判断方案是否具备落地条件。本节结合项目实际依赖、部署方式、页面交互和源码实现，对系统可行性进行分析。",
    )
    add_heading(document, heading_22, "2.1 可行性分析", 2)
    add_heading(document, heading_22, "2.1.1 经济可行性", 3)
    add_body(
        document,
        heading_22,
        "系统所用 Vue 3、Element Plus、Spring Boot、MyBatis-Plus、MySQL、OpenCSV 和 ECharts 均可在开源许可范围内使用，开发与测试可在普通个人电脑上完成，不依赖收费云平台或专用硬件。部署阶段只需准备 Java 运行环境、Web 静态资源和 MySQL 数据库，软件与硬件投入较低，因此具备经济可行性。",
    )
    add_heading(document, heading_22, "2.1.2 操作可行性", 3)
    add_body(
        document,
        heading_22,
        "系统采用浏览器访问方式，普通用户登录后可通过侧边导航进入账户、分类、流水、预算、周期账单、转账、统计、导入和设置页面。各页面主要使用表单、按钮、表格、对话框、进度条和图表完成操作，成功与失败结果通过统一消息提示反馈；管理员入口按角色显示，功能范围清晰。整体交互符合常见 Web 应用使用习惯，普通用户无需掌握专业财务或数据库知识，因此具备操作可行性。",
    )
    add_heading(document, heading_22, "2.1.3 技术可行性", 3)
    add_body(
        document,
        heading_22,
        "前端技术栈能够支持组件化页面、响应式状态、路由权限、文件上传和图表展示；Spring Boot 能够提供 REST 接口、依赖注入、参数校验、拦截器和事务管理；MyBatis-Plus 能够完成分页、条件查询和数据持久化；MySQL 的约束、索引与事务能够支撑七张核心业务表。项目源码、依赖配置、数据库结构和实际运行页面均可相互验证，说明所选技术可以实现系统功能。",
    )
    add_body(
        document,
        heading_22,
        "综上，系统在开发成本、用户操作和技术实现方面均不存在明显障碍，采用的软硬件环境与框架能够支撑后续需求分析、系统设计、功能实现和测试工作。",
    )


def build_section_2(document, figures):
    heading_22 = legacy.find_paragraph(document, "2.2 系统需求分析")
    heading_23 = legacy.find_paragraph(document, "2.3 本章小结")
    legacy.remove_between(heading_22, heading_23)

    add_heading(document, heading_23, "2.2.1 业务需求分析", 3)
    add_body(
        document,
        heading_23,
        "个人财务记账与分析系统面向普通用户和管理员两类角色。普通用户完成账户、收支、预算、周期账单、统计和数据导入等个人财务活动；管理员只负责系统用户列表、用户删除和角色切换。系统后端统一完成 JWT 身份校验、参数校验、业务处理和 MySQL 数据读写，前端负责表单交互、列表展示、ECharts 图表和结果提示。",
    )
    add_body(
        document,
        heading_23,
        "顶层数据流图采用与参考论文一致的数据流图符号：外部实体使用矩形，系统处理过程使用椭圆，数据流使用实线箭头。普通用户向系统提交注册登录、账户、流水、预算、账单、统计查询和 CSV 文件，系统返回余额、列表、预算预警、图表和导入结果；管理员提交用户管理请求并接收处理结果，如图 2-1 所示。",
    )
    add_figure(document, heading_23, figures["top_dfd"], "图 2-1 系统顶层数据流图", 6.0)
    add_page_break(document, heading_23)
    add_body(
        document,
        heading_23,
        "在顶层数据流图基础上，将系统分解为身份认证与用户管理、账户与收支处理、预算与周期计划、统计分析与 CSV 导入四个处理过程，并使用开放式数据存储符号表示 D1 至 D7 七张核心表。为避免多条长线交叉，图中对统计处理所读取的 D3、D4、D5 数据存储进行重复绘制；相同编号始终表示同一张逻辑数据表，如图 2-2 所示。",
    )
    add_figure(document, heading_23, figures["level_one_dfd"], "图 2-2 系统一层数据流图", 6.0)
    add_body(
        document,
        heading_23,
        "业务处理以当前登录用户 ID 作为数据隔离条件。账户、流水、预算、周期账单和预算预警均按 user_id 查询；分类表为系统预置只读数据。账户转账在同一事务内写入一条转出流水和一条转入流水，周期账单生成在同一事务内写入流水并推进下次到期日，保证跨表处理的一致性。",
    )

    add_heading(document, heading_23, "2.2.2 用户需求分析", 3)
    add_body(
        document,
        heading_23,
        "普通用户需要完成注册登录、账户新增与维护、余额汇总、多币种 CNY 参考折算、分类浏览、收支记录增删改查、分页筛选、账户转账、预算设置与预警、周期账单维护与一键生成、月度和年度统计、分类和趋势分析、CSV 导入以及修改密码。图中参与者与主用例之间使用实线无箭头关联，主用例包含的子功能使用带空心箭头的虚线 <<include>> 关系，如图 2-3 所示。",
    )
    add_figure(document, heading_23, figures["user_use_case"], "图 2-3 普通用户用例图", 5.7)
    add_page_break(document, heading_23)
    add_body(
        document,
        heading_23,
        "管理员功能以实际 AdminPage、AdminController 和 AdminServiceImpl 为依据，仅包括管理员登录、查看用户列表、删除用户和切换用户角色。系统没有实现独立的用户名搜索、用户详情或权限配置页面，因此管理员用例图不绘制这些不存在的功能。删除和角色切换作为用户维护的包含用例，所有关联线分区布置且无交叉，如图 2-4 所示。",
    )
    add_figure(document, heading_23, figures["admin_use_case"], "图 2-4 管理员用例图", 5.8)

    add_heading(document, heading_23, "2.2.3 性能与非功能需求分析", 3)
    add_body(
        document,
        heading_23,
        "系统面向个人记账场景，非功能需求重点是响应稳定、金额准确、数据隔离和关键事务一致性。系统应在普通个人电脑和本地网络环境下保持常用页面秒级响应，并对文件大小、单次导入条数、金额精度和访问权限设置明确边界。具体要求如表 2-1 所示。",
    )
    add_table_caption(document, heading_23, "表 2-1 系统非功能需求")
    add_three_line_table_before(
        document,
        heading_23,
        ("需求类别", "具体要求", "实现依据"),
        [
            ("响应性能", "常用查询和页面切换在本地环境保持秒级响应", "分页查询、组合索引、前端并行加载"),
            ("数据准确性", "金额保留两位小数，不使用二进制浮点计算", "MySQL DECIMAL(12,2) 与 Java BigDecimal"),
            ("安全性", "未登录不能访问业务接口，普通用户不能访问管理员接口", "Vue Router、JWT、LoginInterceptor、AdminInterceptor"),
            ("可靠性", "转账和周期账单生成不能出现半成功数据", "@Transactional 事务与统一异常处理"),
            ("容量边界", "CSV 文件不超过 5 MB，单次最多导入 1000 条有效记录", "MultipartFile 大小校验与导入计数限制"),
            ("可维护性", "前后端模块边界清晰，业务规则集中在 Service 层", "Controller-Service-Mapper-Entity 分层"),
        ],
        (1400, 3650, 3250),
        font_size=9,
        center_columns=(0,),
    )

    summary_node = heading_23._p.getnext()
    if summary_node is not None and summary_node.tag == qn("w:p"):
        summary = next(p for p in document.paragraphs if p._p is summary_node)
        summary.clear()
        summary.add_run(
            "本章从技术、经济和操作层面论证了系统可行性，并以实际源码和运行功能为依据完成业务需求、用户需求和非功能需求分析。顶层数据流图与一层数据流图明确了两类外部角色、四类处理过程和七类数据存储；普通用户与管理员用例图进一步界定了功能边界，其中管理员仅具有用户列表、删除和角色切换能力。"
        )
        legacy.format_body_paragraph(summary)


def build_section_3(document):
    heading_3 = legacy.find_paragraph(document, "3 系统相关理论和技术")
    heading_4 = legacy.find_paragraph(document, "4 系统设计")
    legacy.remove_between(heading_3, heading_4)

    add_body(
        document,
        heading_4,
        "本章结合项目配置文件、依赖清单和运行环境，说明系统开发工具、前后端框架、数据持久化、身份认证和可视化技术，并解释各技术在个人财务业务中的具体作用。",
    )
    add_heading(document, heading_4, "3.1 开发与运行环境", 2)
    add_body(
        document,
        heading_4,
        "系统采用前后端分离方式开发，前端工程、后端工程和数据库脚本分别维护。开发环境既要满足前端依赖管理和生产构建，也要支持 Java 服务编译、自动化测试及 MySQL 数据持久化。",
    )
    add_heading(document, heading_4, "3.1.1 前端开发环境", 3)
    add_body(
        document,
        heading_4,
        "前端使用 JavaScript 编写单文件组件和请求逻辑，使用 pnpm 管理依赖，使用 Vite 8 提供开发服务器与生产构建。开发服务器默认运行在 5173 端口，并将 API 请求代理到 8080 端口的后端服务，从而在开发阶段保持页面资源与业务接口相互独立。Visual Studio Code 可用于组件、样式和配置文件编辑。",
    )
    add_heading(document, heading_4, "3.1.2 后端与数据库环境", 3)
    add_body(
        document,
        heading_4,
        "后端使用 Java 21 和 Maven 组织 Spring Boot 工程，内嵌 Tomcat 在 8080 端口提供 REST 接口；MySQL 在 3306 端口保存业务数据。初始化脚本负责创建数据库、七张核心表、约束和索引，实际核验环境的数据库版本为 MySQL 8.4.8。系统开发与运行环境汇总如表3-1所示。",
    )
    add_table_caption(document, heading_4, "表 3-1 系统开发与运行环境")
    add_three_line_table_before(
        document,
        heading_4,
        ("类别", "环境或工具", "用途"),
        [
            ("操作系统", "Windows 10/11", "本地开发、运行与测试"),
            ("开发语言", "Java 21、JavaScript、SQL", "后端业务、前端交互和数据定义"),
            ("开发工具", "IntelliJ IDEA、Visual Studio Code", "后端与前端代码开发"),
            ("构建工具", "Maven 3.9、pnpm、Vite 8", "依赖管理、测试和生产构建"),
            ("运行端口", "前端 5173、后端 8080、MySQL 3306", "页面、API 与数据库访问"),
        ],
        (1350, 3400, 3550),
        font_size=9,
        center_columns=(0,),
    )

    add_heading(document, heading_4, "3.2 开发框架与关键技术", 2)
    add_body(
        document,
        heading_4,
        "系统围绕页面交互、接口处理、数据持久化、访问控制和统计展示选择技术。各框架不只作为依赖引入，还分别承担明确的系统职责，主要技术与版本如表3-2所示。",
    )
    add_heading(document, heading_4, "3.2.1 Vue 3 与前端生态", 3)
    add_body(
        document,
        heading_4,
        "Vue 3 通过组件和响应式状态组织登录、仪表盘、账户、分类、流水、预算、周期账单、转账、统计、导入、设置和管理员 12 个路由页面。Vue Router 负责页面导航与登录守卫，Pinia 保存访问令牌、用户信息和角色状态，Element Plus 提供表单、表格、对话框、上传、分页和进度条组件，Axios 负责统一请求头、响应解包和错误提示。",
    )
    add_heading(document, heading_4, "3.2.2 Spring Boot 与 MyBatis-Plus", 3)
    add_body(
        document,
        heading_4,
        "Spring Boot 负责 Web 服务启动、依赖注入、参数校验、拦截器注册、定时任务和事务管理。后端按 Controller、Service、Mapper、Entity 分层：Controller 接收并校验请求，Service 实现数据隔离、金额规则、预算预警、周期日期推进和转账事务，Mapper 使用 MyBatis-Plus 完成条件查询、分页和持久化，Entity 与数据库字段建立映射。该分层使接口交互、业务规则和数据访问职责相互分离。",
    )
    add_heading(document, heading_4, "3.2.3 MySQL 与事务处理", 3)
    add_body(
        document,
        heading_4,
        "MySQL 保存 user、account、category、transaction、budget、recurring_bill 和 budget_alert 七张核心表。金额字段使用 DECIMAL(12,2)，后端使用 BigDecimal 计算，避免二进制浮点数造成金额误差；账户、流水、预算、周期账单和预警查询均附加 user_id 条件，防止不同用户数据混用。账户转账和周期账单生成使用 @Transactional，在多条记录任一处理失败时整体回滚。",
    )
    add_heading(document, heading_4, "3.2.4 身份认证与统一异常处理", 3)
    add_body(
        document,
        heading_4,
        "用户登录成功后，JJWT 根据用户标识和角色生成访问令牌，前端将令牌加入 Authorization 请求头；LoginInterceptor 校验登录状态，AdminInterceptor 进一步校验管理员角色。用户密码使用 BCrypt 哈希后保存，不在数据库中存储明文。后端统一使用 Result<T> 返回 code、message 和 data，GlobalExceptionHandler 将业务异常、参数异常和系统异常转换为一致的响应结构。",
    )
    add_heading(document, heading_4, "3.2.5 文件导入、图表与汇率服务", 3)
    add_body(
        document,
        heading_4,
        "OpenCSV 负责读取上传文件，系统对文件大小、有效记录数量、日期、类型、金额、账户和分类逐行校验，并返回成功数、失败数和失败原因。ECharts 将后端统计结果转换为分类饼图、月度趋势图和年度对比图。账户支持 CNY、USD、EUR、JPY、GBP、HKD 六种币种，ExchangeRateServiceImpl 提供静态参考汇率及反向汇率，余额汇总时计算 CNY 参考等值；该功能不作为实时结算汇率使用。",
    )
    add_page_break(document, heading_4)
    add_table_caption(document, heading_4, "表 3-2 系统主要技术与版本")
    add_three_line_table_before(
        document,
        heading_4,
        ("层次", "技术", "版本", "主要作用"),
        [
            ("前端", "Vue", "3.5.34", "组件化页面与响应式状态"),
            ("前端", "Vue Router", "5.0.6", "路由组织与权限守卫"),
            ("前端", "Element Plus", "2.13.7", "表单、表格、上传和提示组件"),
            ("前端", "Pinia、Axios、ECharts", "3.0.4、1.15.2、5.6.0", "状态管理、HTTP 请求和图表"),
            ("后端", "Spring Boot", "3.5.14", "Web 服务、依赖注入与事务"),
            ("后端", "MyBatis-Plus", "3.5.15", "对象关系映射与分页查询"),
            ("后端", "JJWT、BCrypt", "0.13.0、Spring Crypto", "令牌认证与密码哈希"),
            ("后端", "OpenCSV", "5.9", "CSV 文件解析"),
            ("数据库", "MySQL", "8.4.x", "业务数据持久化"),
        ],
        (1050, 2100, 1900, 3250),
        font_size=8.7,
        center_columns=(0, 2),
        keep_together=True,
    )
    add_body(
        document,
        heading_4,
        "上述技术共同形成从浏览器页面、HTTP 请求、权限校验、业务处理到 MySQL 持久化的完整链路。Axios 响应拦截器在 code=200 时返回业务数据，在 401 时清理登录状态并跳转登录页，在其他业务错误时显示后端 message，使页面交互与后端统一响应保持一致。",
    )

    add_heading(document, heading_4, "3.3 本章小结", 2)
    add_body(
        document,
        heading_4,
        "本章说明了系统开发环境、前端框架、后端分层、数据库、认证、文件导入和可视化技术。所选技术均已在实际项目配置和运行页面中得到验证，能够支撑用户认证、个人财务数据维护、多币种账户、预算预警、周期账单、统计分析和管理员管理等功能。",
    )


def build_section_4(document, figures):
    heading_4 = legacy.find_paragraph(document, "4 系统设计")
    heading_5 = legacy.find_paragraph(document, "5 系统实现")
    legacy.remove_between(heading_4, heading_5)

    add_body(
        document,
        heading_5,
        "系统设计承接需求分析并为功能实现提供结构依据。本章从总体架构、功能模块、核心流程、数据库和详细模块五个方面展开，所有实体、字段、接口和处理步骤均与当前源码、根目录 sql/01-init.sql 及实际运行数据库相互校核。",
    )
    add_heading(document, heading_5, "4.1 系统总体架构设计", 2)
    add_body(
        document,
        heading_5,
        "系统采用 B/S 架构和前后端分离模式，按表现层、应用层、服务层、数据层和基础架构层组织。表现层为 Vue Web 界面；应用层对应 11 个普通业务页面和管理员页面；服务层封装用户、账户、分类、流水、预算、账单、统计、汇率和管理员业务；数据层承担事务、对象映射、MySQL 持久化和静态汇率参考；基础架构由 Vue 3、Element Plus、Spring Boot、JWT/BCrypt 与 MySQL 8.4 支撑。严格按照参考图格式，各层仅使用独立矩形区域表示，不绘制层与层之间的连接线，如图 4-1 所示。",
    )
    add_figure(document, heading_5, figures["architecture"], "图 4-1 系统总体架构图", 6.1)

    add_heading(document, heading_5, "4.2 系统总体设计", 2)
    add_body(
        document,
        heading_5,
        "系统总体功能按用户认证、账户管理、收支管理、预算管理、周期账单、统计分析和数据及用户管理七个模块划分。模块树使用矩形表示模块，采用横平竖直、无箭头的层级连线表示系统、一级模块和子功能之间的从属关系；图形内统一使用名词、动宾短语或主谓短语，不使用 CRUD、斜杠缩写等歧义表达，并与业务流程图和数据流图的符号严格区分，如图4-2所示。",
    )
    add_figure(document, heading_5, figures["module_tree"], "图 4-2 系统总体功能模块图", 6.1)

    add_heading(document, heading_5, "4.3 功能模块流程设计", 2)
    add_body(
        document,
        heading_5,
        "流程图统一采用参考论文中的标准符号：开始和结束使用圆角终止符，处理步骤使用矩形，用户输入或文件输入使用平行四边形，条件判断使用菱形，执行方向使用实心箭头。所有回退支路沿图形外侧布置，不穿过其他节点。",
    )

    add_heading(document, heading_5, "4.3.1 账户管理流程设计", 3)
    add_body(
        document,
        heading_5,
        "用户进入账户页面后，可新增、编辑或停用现金、银行卡、支付宝和微信账户。前端先校验名称、类型、金额和币种格式，后端再校验账户归属和状态。处理成功后，系统根据初始余额与收支流水计算当前余额，并使用静态参考汇率显示非 CNY 账户的 CNY 参考等值，最后刷新账户列表，流程如图 4-3 所示。",
    )
    add_figure(document, heading_5, figures["flow_account"], "图 4-3 账户管理流程图", 5.5)

    add_heading(document, heading_5, "4.3.2 收支记录与转账流程设计", 3)
    add_body(
        document,
        heading_5,
        "普通收支和账户转账共用交易数据模型。普通收支在校验账户、分类、类型、金额和时间后写入一条 transaction 记录；转账在校验两个账户均属于当前用户且状态有效后，以 @Transactional 事务写入一条支出和一条收入记录，并使用同一 transfer_id 关联。两条记录任一写入失败时整体回滚，流程如图 4-4 所示。",
    )
    add_figure(document, heading_5, figures["flow_transaction"], "图 4-4 收支记录与转账流程图", 5.6)

    add_heading(document, heading_5, "4.3.3 预算管理流程设计", 3)
    add_body(
        document,
        heading_5,
        "预算按用户、支出分类和月份设置。系统通过 user_id、category_id、month 联合唯一约束判断更新或新增预算，再汇总当月同分类支出，计算消耗百分比及 NORMAL、DAILY_WARN、MONTHLY_WARN、OVERSPENT 四级预警。BudgetScheduler 每日凌晨 2 点重新计算并持久化预警，流程如图 4-5 所示。",
    )
    add_figure(document, heading_5, figures["flow_budget"], "图 4-5 预算管理流程图", 5.5)

    add_heading(document, heading_5, "4.3.4 周期账单生成流程设计", 3)
    add_body(
        document,
        heading_5,
        "周期账单支持 daily、weekly、monthly、yearly 四种周期。用户维护名称、账户、分类、金额、类型、周期和下次到期日；一键生成时，服务层在同一事务中写入收支流水，并循环推进 next_due_date，直至该日期晚于当前日期。该规则避免历史到期模板生成后仍停留在过去，流程如图 4-6 所示。",
    )
    add_figure(document, heading_5, figures["flow_recurring"], "图 4-6 周期账单生成流程图", 5.5)

    add_heading(document, heading_5, "4.3.5 CSV 导入流程设计", 3)
    add_body(
        document,
        heading_5,
        "CSV 导入先校验账户归属、文件扩展名和 5 MB 大小上限，再由 OpenCSV 逐行读取日期、分类 ID、类型、金额和备注。每行校验列数、分类、类型、金额和时间；有效行立即写入流水并累计成功数，非法行记录行号和失败原因后继续处理，最多接收 1000 条有效记录。所有循环线均从节点外侧回到“读取下一行”，不存在穿框或交叉连线，如图 4-7 所示。",
    )
    add_figure(document, heading_5, figures["flow_csv"], "图 4-7 CSV 导入流程图", 5.4)

    add_heading(document, heading_5, "4.4 系统数据库设计", 2)
    add_heading(document, heading_5, "4.4.1 数据库概念结构设计", 3)
    add_body(
        document,
        heading_5,
        "数据库概念结构围绕用户财务活动设计，包含用户、账户、收支分类、收支记录、预算、周期账单和预算预警七个核心实体。实体使用矩形表示，属性使用椭圆表示，实体与属性之间使用无箭头实线连接，主键属性在椭圆内加下划线，符号与参考论文保持一致。",
    )
    entity_sections = [
        (
            "用户实体保存登录账号、BCrypt 密码哈希、角色和审计时间。用户是账户、流水、预算和周期账单的归属主体，用户实体属性如图 4-8 所示。",
            "attr_user",
            "图 4-8 用户实体属性图",
        ),
        (
            "账户实体保存账户名称、类型、初始余额、币种、状态和审计时间。当前余额由初始余额与关联流水实时计算，不在表中重复存储，账户实体属性如图 4-9 所示。",
            "attr_account",
            "图 4-9 账户实体属性图",
        ),
        (
            "收支分类实体保存分类名称和收入/支出类型，系统初始化 8 个支出分类和 5 个收入分类，普通用户只读浏览，收支分类实体属性如图 4-10 所示。",
            "attr_category",
            "图 4-10 收支分类实体属性图",
        ),
        (
            "收支记录实体保存用户、账户、分类、交易类型、金额、时间和备注。转账产生的两条流水通过 transfer_id 建立逻辑自关联，收支记录实体属性如图 4-11 所示。",
            "attr_transaction",
            "图 4-11 收支记录实体属性图",
        ),
        (
            "预算实体按用户、支出分类和月份保存预算金额，并为预算进度和预警计算提供上限值，预算实体属性如图 4-12 所示。",
            "attr_budget",
            "图 4-12 预算实体属性图",
        ),
        (
            "周期账单实体保存固定收入或支出的模板、周期、状态和下次到期日，并关联账户与分类，周期账单实体属性如图 4-13 所示。",
            "attr_recurring",
            "图 4-13 周期账单实体属性图",
        ),
        (
            "预算预警实体保存定时计算得到的月份、预警级别、预算金额、已支出金额和消耗百分比。该表已在实际根目录 SQL、实体类和 MySQL 库中实现，预算预警实体属性如图 4-14 所示。",
            "attr_alert",
            "图 4-14 预算预警实体属性图",
        ),
    ]
    for index, (intro, figure_key, caption) in enumerate(entity_sections):
        if index in (0, 2, 4, 6):
            add_page_break(document, heading_5)
        add_body(document, heading_5, intro)
        add_figure(document, heading_5, figures[figure_key], caption, 5.8)

    add_page_break(document, heading_5)
    add_body(
        document,
        heading_5,
        "系统总体 E-R 图采用 Chen 表示法：实体为矩形，联系为菱形，连接线不带箭头，线端标注 1 和 N。用户拥有账户并记录流水、设置预算和维护周期账单；账户承载流水并关联周期账单；分类用于归类流水、预算、账单和预警；预算可生成多次预警记录。实体在同一图中只出现一次，图中不存在重复实体或交叉连线，如图 4-15 所示。",
    )
    add_figure(document, heading_5, figures["er"], "图 4-15 系统总体 E-R 图", 6.1)

    add_heading(document, heading_5, "4.4.2 数据库逻辑结构设计", 3)
    add_body(
        document,
        heading_5,
        "根据概念模型，将七个实体转换为关系模式。下列关系模式中带下划线的字段为主键；系统未声明物理外键，而是在 Service 层校验 user_id、account_id 和 category_id 的归属与状态，并在管理员删除用户时按子表到父表顺序执行级联删除。",
    )
    legacy.add_relation_before(document, heading_5, "用户 user", "id", "username，password，role，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "账户 account", "id", "user_id，name，type，initial_balance，currency，status，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "分类 category", "id", "name，type，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "收支记录 transaction", "id", "user_id，account_id，category_id，type，amount，note，time，transfer_id，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "预算 budget", "id", "user_id，category_id，month，amount，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "周期账单 recurring_bill", "id", "user_id，account_id，category_id，name，amount，type，period，next_due_date，status，create_time，update_time")
    legacy.add_relation_before(document, heading_5, "预算预警 budget_alert", "id", "user_id，category_id，month，alert_level，budget_amount，spent_amount，percentage，create_time，update_time")

    add_heading(document, heading_5, "4.4.3 数据库物理结构设计", 3)
    add_body(
        document,
        heading_5,
        "数据库采用 MySQL 8.4、InnoDB 引擎和 utf8mb4 字符集。金额统一使用 DECIMAL，Java 端使用 BigDecimal；日期时间使用 DATE 或 DATETIME；查询热点字段建立普通索引或联合索引。各表按参考论文的三线表格式列出。",
    )
    for table_index, (table_name, chinese_name, intro, rows) in enumerate(legacy.database_tables(), 1):
        intro_paragraph = add_body(document, heading_5, intro)
        intro_paragraph.paragraph_format.keep_with_next = True
        add_table_caption(document, heading_5, f"表 4-{table_index} {table_name} {chinese_name}")
        add_three_line_table_before(
            document,
            heading_5,
            ("字段名", "字段描述", "数据类型", "完整性约束"),
            rows,
            (1450, 2900, 1700, 2300),
            font_size=8.5,
            center_columns=(0, 2),
            keep_together=True,
        )

    add_heading(document, heading_5, "4.5 详细模块设计", 2)
    add_heading(document, heading_5, "4.5.1 核心实体类设计", 3)
    add_body(
        document,
        heading_5,
        "后端实体类与数据库表一一映射，并由 Lombok 生成访问方法。业务操作不直接写入实体类，而由对应 Service 完成，因此类图在实体属性区列出关键字段，在方法区列出由业务服务提供的主要操作。用户、账户、分类、流水、预算、周期账单和预算预警之间的多重关系与 E-R 图保持一致，如图 4-16 所示。",
    )
    add_figure(document, heading_5, figures["class"], "图 4-16 系统核心类图", 6.1)

    add_page_break(document, heading_5)
    add_heading(document, heading_5, "4.5.2 核心业务时序设计", 3)
    sequences = [
        (
            "（1）登录时序。用户在 LoginPage 输入用户名和密码，Axios 调用 UserController，UserService 通过 UserMapper 查询 user 表，使用 BCrypt 校验密码后生成 JWT，并将 token、用户 ID、用户名和角色返回前端，如图 4-17 所示。",
            "seq_login",
            "图 4-17 登录核心时序图",
        ),
        (
            "（2）转账时序。TransactionService 在事务中校验账户归属与余额，向 transaction 表写入转出和转入两条记录，并使用同一 transfer_id 关联，前端收到结果后刷新余额和流水，如图 4-18 所示。",
            "seq_transfer",
            "图 4-18 账户转账核心时序图",
        ),
        (
            "（3）预算时序。BudgetService 查询同月同分类预算并执行新增或更新，再通过 TransactionMapper 汇总实际支出，计算进度与预警级别后返回 BudgetPage，如图 4-19 所示。",
            "seq_budget",
            "图 4-19 预算管理核心时序图",
        ),
        (
            "（4）周期账单时序。RecurringBillService 查询账单及关联账户、分类，在事务中写入流水并推进 next_due_date，最终刷新账单和统计结果，如图 4-20 所示。",
            "seq_recurring",
            "图 4-20 周期账单生成核心时序图",
        ),
        (
            "（5）CSV 导入时序。TransactionService 使用 OpenCSV 读取每一行，校验后通过 Mapper 写入有效流水，累计成功数、失败数和失败原因并返回 ImportResultDTO，如图 4-21 所示。",
            "seq_csv",
            "图 4-21 CSV 导入核心时序图",
        ),
    ]
    for sequence_index, (text_value, figure_key, caption) in enumerate(sequences):
        if sequence_index:
            add_page_break(document, heading_5)
        add_body(document, heading_5, text_value)
        add_figure(document, heading_5, figures[figure_key], caption, 6.1)

    add_page_break(document, heading_5)
    add_heading(document, heading_5, "4.5.3 核心模块职责设计", 3)
    add_body(
        document,
        heading_5,
        "各业务模块采用 Controller-Service-Mapper-Entity 纵向分层，控制层只负责请求入口、参数绑定和统一响应，业务规则集中在 Service 层，Mapper 负责数据访问。核心类与职责如表 4-8 所示。",
    )
    add_table_caption(document, heading_5, "表 4-8 核心模块类职责")
    add_three_line_table_before(
        document,
        heading_5,
        ("模块", "控制层", "业务层", "数据访问层", "主要职责"),
        [
            ("用户认证", "UserController", "UserServiceImpl", "UserMapper", "注册、登录、改密、JWT 与 BCrypt"),
            ("账户管理", "AccountController", "AccountServiceImpl", "AccountMapper", "账户维护、余额计算、多币种折算"),
            ("分类浏览", "CategoryController", "CategoryServiceImpl", "CategoryMapper", "读取预置收入与支出分类"),
            ("收支管理", "TransactionController", "TransactionServiceImpl", "TransactionMapper", "流水维护、筛选、转账与 CSV 导入"),
            ("预算预警", "BudgetController", "BudgetServiceImpl、AlertService", "BudgetMapper、AlertMapper", "预算保存、进度计算与定时预警"),
            ("周期账单", "RecurringBillController", "RecurringBillServiceImpl", "RecurringBillMapper", "模板维护、一键生成和日期推进"),
            ("统计分析", "StatisticsController", "StatisticsServiceImpl", "TransactionMapper", "月度、年度、分类和趋势聚合"),
            ("汇率参考", "ExchangeRate\nController", "ExchangeRate\nServiceImpl", "无数据库表", "六种币种静态参考汇率"),
            ("用户管理", "AdminController", "AdminServiceImpl", "六个业务 Mapper", "用户列表、级联删除与角色切换"),
        ],
        (900, 1900, 2000, 1650, 1850),
        font_size=7.4,
        center_columns=(0, 1, 2, 3),
        keep_together=True,
    )

    add_heading(document, heading_5, "4.6 本章小结", 2)
    add_body(
        document,
        heading_5,
        "本章完成了系统总体架构、功能模块、五类核心流程、七实体数据库和详细模块设计。架构图取消层间连接线；流程图采用终止符、处理框、输入输出框和判断菱形；数据库补齐七张实体属性图，并以无重复实体、无交叉连线的 Chen E-R 图表达关系；类图和五张时序图进一步说明了前端、控制层、业务层、数据访问层与 MySQL 的协作过程。",
    )


def build_section_5(document, screenshots: Path):
    heading_5 = legacy.find_paragraph(document, "5 系统实现")
    heading_6 = legacy.find_paragraph(document, "6 系统测试")
    legacy.remove_between(heading_5, heading_6)

    def shot(name: str):
        candidate = screenshots / name
        return candidate if candidate.exists() else None

    add_body(
        document,
        heading_6,
        "系统实现章节结合前端页面、Axios API 模块、Controller、ServiceImpl、Mapper 和 MySQL 表说明各功能如何落地。每节均给出模块职责、实现思路、关键接口、涉及数据表和运行效果，截图均来自本机实际系统页面。",
    )

    add_heading(document, heading_6, "5.1 登录注册模块", 2)
    add_body(
        document,
        heading_6,
        "登录注册模块由 LoginPage.vue、api/user.js、api/request.js、UserController、UserServiceImpl、UserMapper 和 user 表共同实现。前端使用 Element Plus 的 el-form、el-tabs、el-input 与 el-button 组织登录和注册表单，提交时通过 Axios 调用 /api/v1/user/register 与 /api/v1/user/login。",
    )
    add_body(
        document,
        heading_6,
        "后端注册流程先校验用户名唯一性，再使用 BCryptPasswordEncoder 生成密码哈希后写入 user 表；登录流程按用户名查询用户，校验密码后签发 JWT，并返回 token、userId、username 和 role。前端把令牌写入 Pinia 与 localStorage，后续请求由 Axios 拦截器自动附加 Authorization 头。改密接口 POST /api/v1/user/change-password 需要旧密码校验。实现效果为用户可完成注册、登录和密码维护，未登录访问受保护页面时由 Vue Router 与 LoginInterceptor 共同拦截，登录注册页面如图5-1所示。",
    )
    login_shot = shot("login.png")
    if login_shot:
        add_figure(document, heading_6, login_shot, "图 5-1 登录注册页面", 5.8)

    add_heading(document, heading_6, "5.2 用户管理模块", 2)
    add_body(
        document,
        heading_6,
        "用户管理模块对应管理员能力，由 AdminPage.vue、api/admin.js、AdminController、AdminServiceImpl、AdminInterceptor 和 user 表实现。前端路由将 AdminPage 标记为 requiresAdmin，普通用户不能进入；后端管理员接口同样经过 AdminInterceptor 校验角色。",
    )
    add_body(
        document,
        heading_6,
        "模块提供用户列表查询、删除用户和角色切换，接口包括 GET /api/v1/admin/users、DELETE /api/v1/admin/users/{userId} 和 PUT /api/v1/admin/users/{userId}/role。删除用户时 Service 按预算预警、周期账单、预算、流水、账户到用户的顺序清理关联数据，避免残留孤儿记录。实现效果为管理员可查看系统用户并维护异常账号，系统不提供独立的用户搜索或详情页，管理员用户管理页面如图5-2所示。",
    )
    admin_shot = shot("admin.png")
    if admin_shot:
        add_figure(document, heading_6, admin_shot, "图 5-2 管理员用户管理页面", 5.8)

    add_heading(document, heading_6, "5.3 收支记账模块", 2)
    add_body(
        document,
        heading_6,
        "收支记账是系统核心业务，由 TransactionListPage.vue、TransferPage.vue、api/transaction.js、TransactionController、TransactionServiceImpl、TransactionMapper、AccountServiceImpl 以及 transaction、account、category 表实现。普通记账在校验账户、分类、类型、金额和时间后写入一条流水；账户转账在同一事务中写入支出和收入两条记录，并用 transfer_id 关联。",
    )
    add_body(
        document,
        heading_6,
        "系统支持流水分页查询、多条件筛选、修改、删除和账户转账。主要接口为 GET /api/v1/transaction、POST /api/v1/transaction、PUT /api/v1/transaction/{id}、DELETE /api/v1/transaction/{id} 与 POST /api/v1/transaction/transfer。所有查询均以当前登录用户 ID 做数据隔离。实现效果是用户可以维护日常收支，并在不同账户之间安全转账；收支记录列表页面如图5-3所示，账户转账页面如图5-4所示。",
    )
    transaction_shot = shot("transaction.png")
    transfer_shot = shot("transfer.png")
    if transaction_shot:
        add_figure(document, heading_6, transaction_shot, "图 5-3 收支记录列表页面", 5.8)
    if transfer_shot:
        add_figure(document, heading_6, transfer_shot, "图 5-4 账户转账页面", 5.8)

    add_heading(document, heading_6, "5.4 收支分类模块", 2)
    add_body(
        document,
        heading_6,
        "收支分类模块由 CategoryPage.vue、api/category.js、CategoryController、CategoryServiceImpl、CategoryMapper 和 category 表实现。分类为系统初始化种子数据，包含餐饮、交通、住房、娱乐、医疗、教育、购物、其他等支出分类，以及工资、奖金、兼职、理财、其他等收入分类。普通用户只读浏览，不提供增删改。",
    )
    add_body(
        document,
        heading_6,
        "列表接口为 GET /api/v1/category，分类统计由 GET /api/v1/statistics/category-summary 提供。前端 CategoryPage 使用 el-tabs、el-table 和 el-tag 分栏展示收入与支出分类，并可结合分类汇总金额展示。实现效果是记账、预算和统计始终基于统一分类字典，收支分类页面如图5-5所示。",
    )
    category_shot = shot("category.png")
    if category_shot:
        add_figure(document, heading_6, category_shot, "图 5-5 收支分类页面", 5.8)

    add_heading(document, heading_6, "5.5 预算管理模块", 2)
    add_body(
        document,
        heading_6,
        "预算管理模块由 BudgetPage.vue、api/budget.js、BudgetController、BudgetServiceImpl、BudgetAlertServiceImpl、BudgetScheduler 以及 budget、budget_alert、category、transaction 表实现。用户可按月份和支出分类设置预算金额，系统通过 user_id、category_id、month 联合唯一约束保证同一预算只保留一条记录。",
    )
    add_body(
        document,
        heading_6,
        "模块接口包括 GET /api/v1/budget、POST /api/v1/budget、DELETE /api/v1/budget/{id}、GET /api/v1/budget/progress 和 GET /api/v1/budget/alert。预算进度根据 transaction 表中的实际支出计算，前端使用 el-progress 展示消耗比例，并对接近超支和超支状态给出提示。BudgetScheduler 每日凌晨重新计算并持久化预警。实现效果是用户可及时感知预算使用情况，预算管理页面如图5-6所示。",
    )
    budget_shot = shot("budget.png")
    if budget_shot:
        add_figure(document, heading_6, budget_shot, "图 5-6 预算管理页面", 5.8)

    add_heading(document, heading_6, "5.6 财务统计分析模块", 2)
    add_body(
        document,
        heading_6,
        "财务统计分析模块由 DashboardPage.vue、AnalyticsPage.vue、api/statistics.js、StatisticsController、StatisticsServiceImpl 以及 transaction、category、budget 等表实现。后端按月份、年度、分类和趋势等维度聚合交易数据，前端通过 ECharts 渲染为饼图、折线图和预算对比图。",
    )
    add_body(
        document,
        heading_6,
        "模块接口包括 GET /api/v1/statistics/monthly、GET /api/v1/statistics/yearly、GET /api/v1/statistics/category-summary 和 GET /api/v1/statistics/trend。实现效果是用户可查看本月收入、本月支出、结余、分类占比和月度趋势，从流水明细上升为财务概览；仪表盘统计页面如图5-7所示，财务分析页面如图5-8所示。",
    )
    dashboard_shot = shot("dashboard.png")
    analytics_shot = shot("analytics.png")
    if dashboard_shot:
        add_figure(document, heading_6, dashboard_shot, "图 5-7 仪表盘统计页面", 5.8)
    if analytics_shot:
        add_figure(document, heading_6, analytics_shot, "图 5-8 财务分析页面", 5.8)

    add_heading(document, heading_6, "5.7 个人信息与数据扩展模块", 2)
    add_body(
        document,
        heading_6,
        "个人信息管理由 UserSettingsPage.vue、stores/user.js、utils/jwt.js、api/user.js、UserController 和 user 表实现。前端从 Pinia 用户状态和 JWT 载荷读取当前用户信息，在个人设置页展示用户名和角色，并提供修改密码入口。改密接口为 POST /api/v1/user/change-password，后端校验旧密码后写入新的 BCrypt 哈希，个人设置页面如图5-9所示。",
    )
    add_body(
        document,
        heading_6,
        "数据扩展能力由账户管理、周期账单和 CSV 导入共同构成。账户模块通过 AccountPage.vue 与 AccountServiceImpl 维护多账户余额，并基于静态参考汇率显示非 CNY 账户的 CNY 等值；周期账单模块通过 RecurringBillPage.vue 与 RecurringBillServiceImpl 维护模板并一键生成流水；CSV 导入模块通过 ImportPage.vue 与 TransactionServiceImpl.importCsv 完成批量记账。实现效果是用户既能维护个人资料，也能批量扩展历史数据。账户管理页面如图5-10所示，周期账单页面如图5-11所示，CSV 导入页面如图5-12所示。",
    )
    settings_shot = shot("settings.png")
    account_shot = shot("account.png")
    recurring_shot = shot("recurring-bill.png")
    import_shot = shot("import.png")
    if settings_shot:
        add_figure(document, heading_6, settings_shot, "图 5-9 个人设置页面", 5.8)
    if account_shot:
        add_figure(document, heading_6, account_shot, "图 5-10 账户管理页面", 5.8)
    if recurring_shot:
        add_figure(document, heading_6, recurring_shot, "图 5-11 周期账单页面", 5.8)
    if import_shot:
        add_figure(document, heading_6, import_shot, "图 5-12 CSV 导入页面", 5.8)

    add_heading(document, heading_6, "5.8 本章小结", 2)
    add_body(
        document,
        heading_6,
        "本章结合实际源码说明了系统各模块的实现过程。系统实现不是简单的页面展示，而是由 Vue 前端组件、Axios API 封装、Spring Boot 控制层、Service 业务层、Mapper 数据访问和 MySQL 数据表协同完成。登录注册、用户管理、收支记账、分类查询、预算管理、统计分析、个人信息、账户管理、周期账单和 CSV 导入均已对应到真实页面、接口和数据表。",
    )


def build_section_6(document):
    heading_6 = legacy.find_paragraph(document, "6 系统测试")
    heading_7 = legacy.find_paragraph(document, "7 总结和展望")
    legacy.remove_between(heading_6, heading_7)

    add_body(
        document,
        heading_7,
        "系统测试用于验证功能是否满足需求、接口返回是否规范、异常路径是否可控、数据是否正确写入和查询，以及前端构建是否通过。测试范围覆盖用户登录、账户管理、收支记录、转账、预算、周期账单、统计分析、CSV 导入和管理员功能。",
    )

    add_heading(document, heading_7, "6.1 测试概念", 2)
    add_body(
        document,
        heading_7,
        "本系统测试采用黑盒功能验证与白盒业务规则验证相结合的方式。黑盒测试从用户操作和接口输入输出出发，检查页面流程和结果提示；白盒测试关注 Service 层事务、金额精度、数据隔离和异常码。测试环境为本地 Windows 开发环境，前端 5173 端口，后端 8080 端口，MySQL 3306 端口。",
    )

    add_heading(document, heading_7, "6.2 测试内容和方法", 2)
    add_body(
        document,
        heading_7,
        "功能测试用于验证各页面和接口是否完成业务主流程，测试时模拟执行新增、修改、查询、删除和统计操作。接口测试用于验证 /api/v1 下各接口的入参、响应码和 Result<T> 返回格式。异常测试用于验证非法金额、未登录访问、跨用户账户等边界条件。事务测试用于验证转账和周期账单一键生成的原子性。构建测试执行 pnpm build，单元测试执行 Maven test。",
    )

    add_heading(document, heading_7, "6.3 测试用例", 2)
    add_body(
        document,
        heading_7,
        "结合实际系统功能，按模块设计测试用例。表中“实际结果”与“是否通过”依据本地页面联调、接口验证和自动化测试填写。用户认证、账户管理、收支转账、预算管理、周期账单与 CSV 导入、统计分析与管理员测试用例分别如表 6-1 至表 6-6 所示。",
    )

    def add_case_table(caption, rows):
        add_table_caption(document, heading_7, caption)
        add_three_line_table_before(
            document,
            heading_7,
            ("测试编号", "测试内容", "操作步骤", "预期结果", "实际结果", "是否通过"),
            rows,
            (900, 1400, 2200, 2200, 1100, 900),
            font_size=8,
            center_columns=(0, 5),
            keep_together=True,
        )

    add_case_table(
        "表 6-1 用户认证测试用例",
        [
            ("TC01", "合法登录", "输入合法用户名和密码后登录", "返回 token 并进入首页", "与预期一致", "通过"),
            ("TC02", "错误密码登录", "输入错误密码后登录", "提示用户名或密码错误", "与预期一致", "通过"),
            ("TC03", "用户注册", "填写合法用户名和密码注册", "注册成功并可登录", "与预期一致", "通过"),
            ("TC04", "修改密码", "输入旧密码和新密码提交", "密码更新成功", "与预期一致", "通过"),
        ],
    )
    add_case_table(
        "表 6-2 账户管理测试用例",
        [
            ("TC05", "新增账户", "新增现金账户并查询列表", "列表出现新账户且状态有效", "与预期一致", "通过"),
            ("TC06", "编辑账户", "修改账户名称后保存", "列表显示更新后的名称", "与预期一致", "通过"),
            ("TC07", "停用有流水账户", "尝试停用存在流水的账户", "系统拒绝或提示处理关联数据", "与预期一致", "通过"),
            ("TC08", "多币种折算", "查看非 CNY 账户余额", "显示原币种金额和 CNY 参考值", "与预期一致", "通过"),
        ],
    )
    add_case_table(
        "表 6-3 收支记录与转账测试用例",
        [
            ("TC09", "新增支出", "新增一条支出记录", "transaction 写入成功并分页展示", "与预期一致", "通过"),
            ("TC10", "筛选流水", "按时间段和分类筛选", "只返回当前用户且满足条件的数据", "与预期一致", "通过"),
            ("TC11", "账户转账", "银行卡转账到现金账户", "生成一收一支且 transfer_id 相同", "与预期一致", "通过"),
            ("TC12", "删除流水", "删除本人的一条流水", "列表不再显示该记录", "与预期一致", "通过"),
        ],
    )
    add_case_table(
        "表 6-4 预算管理测试用例",
        [
            ("TC13", "设置预算", "按月份和分类设置预算金额", "预算保存成功", "与预期一致", "通过"),
            ("TC14", "查看进度", "打开预算进度页面", "显示消耗比例和剩余金额", "与预期一致", "通过"),
            ("TC15", "超支预警", "支出超过预算后查看预警", "显示超支或预警状态", "与预期一致", "通过"),
            ("TC16", "删除预算", "删除一条预算记录", "预算列表更新", "与预期一致", "通过"),
        ],
    )
    add_case_table(
        "表 6-5 周期账单与 CSV 导入测试用例",
        [
            ("TC17", "创建周期账单", "新增月度账单模板", "账单列表显示新模板", "与预期一致", "通过"),
            ("TC18", "一键生成", "对到期账单执行一键生成", "生成流水并推进下次到期日", "与预期一致", "通过"),
            ("TC19", "合法 CSV 导入", "选择账户并上传合法文件", "成功行写入数据库", "与预期一致", "通过"),
            ("TC20", "非法行处理", "上传含错误行的 CSV", "失败行返回原因且不中断整批", "与预期一致", "通过"),
        ],
    )
    add_case_table(
        "表 6-6 统计分析与管理员测试用例",
        [
            ("TC21", "月度统计", "打开仪表盘查看月度数据", "返回收入、支出和结余", "与预期一致", "通过"),
            ("TC22", "分类统计", "查看分类汇总图", "返回分类金额与占比", "与预期一致", "通过"),
            ("TC23", "管理员列表", "管理员查看用户列表", "返回系统用户数据", "与预期一致", "通过"),
            ("TC24", "角色切换", "管理员修改用户角色", "角色更新成功", "与预期一致", "通过"),
        ],
    )

    add_heading(document, heading_7, "6.4 测试分析与结果", 2)
    add_body(
        document,
        heading_7,
        "从测试结果看，系统主要功能能够闭环运行。用户登录、前端路由守卫和后端拦截器能够共同拦截未授权访问；账户、分类和流水数据均按用户隔离；预算与周期账单可在事务中与账户流水协同；统计分析能够从交易流水中聚合结果；CSV 导入能够对非法数据进行逐行校验；管理员页面和接口能够完成用户信息管理。",
    )
    add_body(
        document,
        heading_7,
        "测试重点关注金额精度、转账原子性和异常提示。金额字段采用 DECIMAL 与 BigDecimal 后，可避免 Double 运算的精度问题；转账过程通过事务保证两条流水一致；异常路径通过 BusinessException 与 GlobalExceptionHandler 转换为统一 Result 错误响应，前端使用 ElMessage 统一提示。",
    )
    add_body(
        document,
        heading_7,
        "验证期间前端 pnpm build 构建成功；后端 Maven 共执行 175 项测试，其中 174 项通过。唯一未通过项是 CrossModuleIntegrationTest.generate_nextDueDateAdvances 的固定日期断言：用例仍期望 2026-07-01，而执行日已晚于该日期，服务按设计循环推进到未来日期 2026-08-01。该结果说明业务实现符合“下次到期日必须晚于当前日期”的规则，但测试用例依赖系统时钟，后续应注入 Clock 或使用相对日期消除时间漂移。",
    )

    add_heading(document, heading_7, "6.5 本章小结", 2)
    add_body(
        document,
        heading_7,
        "本章按用户认证、账户管理、收支转账、预算管理、周期账单与导入、统计分析与管理员六个模块设计测试用例，并结合功能测试、接口测试、异常测试、事务测试、构建测试和单元测试验证系统。主要业务功能与关键约束符合预期，同时识别出一项依赖固定日期的测试维护问题。",
    )


def build_section_7(document):
    heading_7 = legacy.find_paragraph(document, "7 总结和展望")
    node = heading_7._p.getnext()
    parent = heading_7._p.getparent()
    while node is not None and node.tag != qn("w:sectPr"):
        nxt = node.getnext()
        parent.remove(node)
        node = nxt

    end_anchor = document.add_paragraph("__END_ANCHOR__")
    try:
        add_heading(document, end_anchor, "7.1 总结", 2)
        add_body(
            document,
            end_anchor,
            "本文围绕实际系统源码完成了个人财务记账与分析系统的分析、设计与实现说明。系统采用 Vue 3、Element Plus、Pinia、Axios 和 ECharts 的前端技术栈，以及 Spring Boot、MyBatis-Plus、MySQL、JWT 和 BCrypt 的后端技术栈，形成了从页面交互、接口调用、业务处理到数据存储的完整链路。",
        )
        add_body(
            document,
            end_anchor,
            "系统实现了普通用户注册登录、账户管理、分类浏览、流水分页与筛选、账户转账、预算管理、周期账单一键生成流水、统计分析、CSV 导入等功能，也实现了管理员用户列表、删除用户和角色切换能力。数据库设计覆盖用户、账户、分类、流水、预算、周期账单和预算预警七个核心对象；后端统一响应格式、异常处理和登录拦截，前端统一路由、组件、状态管理和请求封装。综合来看，系统能够支撑个人财务记账、分析和管理。",
        )
        add_heading(document, end_anchor, "7.2 展望", 2)
        add_body(
            document,
            end_anchor,
            "后续系统仍有四个方向可以继续完善。第一，增强数据分析能力，补充消费预测、预算建议、同比环比分析和更细粒度的图表维度。第二，增强账户能力，支持更多币种、实时汇率、信用卡账单和资产负债视图。第三，增强数据导入能力，适配不同银行和支付平台的账单格式，降低用户手工上传 CSV 的成本。第四，增强安全、部署与测试能力，完善 JWT 密钥管理、数据库备份和容器化部署，为日期相关服务注入 Clock，并补充稳定的回归测试与日志监控。",
        )
    finally:
        end_anchor._element.getparent().remove(end_anchor._element)


def build_document(input_path: Path, output_path: Path, figure_dir: Path, screenshot_dir: Path):
    figures = generate_all_figures(figure_dir)
    document = Document(str(input_path))
    ensure_section_properties(document)

    build_abstract_and_early_sections(document)
    build_section_2(document, figures)
    build_section_3(document)
    build_section_4(document, figures)
    build_section_5(document, screenshot_dir)
    build_section_6(document)
    build_section_7(document)
    format_chapter_headings(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    check = Document(str(output_path))
    texts = "\n".join(paragraph.text for paragraph in check.paragraphs)
    required = [
        "1.1.1 研究背景",
        "1.1.2 研究意义",
        "1.2 主要研究内容",
        "2.1.1 经济可行性",
        "2.1.2 操作可行性",
        "2.1.3 技术可行性",
        "2.2 系统需求分析",
        "图2-1 系统顶层数据流图",
        "图2-2 系统一层数据流图",
        "图2-3 普通用户用例图",
        "图2-4 管理员用例图",
        "3.1 开发与运行环境",
        "3.1.1 前端开发环境",
        "3.1.2 后端与数据库环境",
        "3.2.1 Vue 3 与前端生态",
        "3.2.2 Spring Boot 与 MyBatis-Plus",
        "3.2.3 MySQL 与事务处理",
        "3.2.4 身份认证与统一异常处理",
        "3.2.5 文件导入、图表与汇率服务",
        "4.1 系统总体架构设计",
        "图4-1 系统总体架构图",
        "图4-7 CSV 导入流程图",
        "图4-8 用户实体属性图",
        "图4-15 系统总体 E-R 图",
        "图4-21 CSV 导入核心时序图",
        "5.1 登录注册模块",
        "6.3 测试用例",
        "表6-1 用户认证测试用例",
        "表6-6 统计分析与管理员测试用例",
        "7.1 总结",
        "7.2 展望",
    ]
    missing = [item for item in required if item not in texts]
    if missing:
        raise RuntimeError(f"Missing required content: {missing}")

    bad = []
    for paragraph in check.paragraphs:
        t = paragraph.text.strip()
        if t in (
            "致谢",
            "参考文献",
            "国内外研究现状",
            "论文组织架构",
            "1.3 国内外研究现状",
            "1.4 论文组织架构",
        ):
            bad.append(t)
    if bad:
        raise RuntimeError(f"Forbidden sections present: {bad}")

    return {
        "paragraphs": len(check.paragraphs),
        "tables": len(check.tables),
        "inline_shapes": len(check.inline_shapes),
        "figures": len(figures),
        "output": str(output_path),
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input",
        type=Path,
        default=ROOT
        / "选题标定-第12题-个人财务记账与分析系统"
        / "论文"
        / "个人财务记账与分析系统论文_最终版.docx",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT
        / "选题标定-第12题-个人财务记账与分析系统"
        / "论文"
        / "个人财务记账与分析系统论文_全篇最终版.docx",
    )
    parser.add_argument(
        "--figure-dir",
        type=Path,
        default=ROOT / "work" / "full_paper_revision" / "generated",
    )
    parser.add_argument(
        "--screenshot-dir",
        type=Path,
        default=ROOT / "work" / "full_paper_revision" / "system_screenshots",
    )
    args = parser.parse_args()
    result = build_document(args.input, args.output, args.figure_dir, args.screenshot_dir)
    for key, value in result.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
