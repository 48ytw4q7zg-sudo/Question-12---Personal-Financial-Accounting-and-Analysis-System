import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = Path(r'C:\Users\Administrator\Desktop\Question-12---Personal-Financial-Accounting-and-Analysis-System-master\选题标定-第12题-个人财务记账与分析系统\论文\个人财务记账与分析系统论文_2.2-7.2格式复验优化版.docx')
doc = Document(str(path))

# 1) collect captions and sequence continuity
figs = []
tabs = []
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    m = re.match(r'^图(\d+)-(\d+)\s+(.+)$', t)
    if m:
        figs.append((i,int(m.group(1)),int(m.group(2)),m.group(3),t))
    m = re.match(r'^表(\d+)-(\d+)\s+(.+)$', t)
    if m:
        tabs.append((i,int(m.group(1)),int(m.group(2)),m.group(3),t))

print('=== FIGURE NUMBER CONTINUITY ===')
by_ch = {}
for f in figs:
    by_ch.setdefault(f[1], []).append(f[2])
for ch, nums in sorted(by_ch.items()):
    expect = list(range(1, max(nums)+1))
    missing = [n for n in expect if n not in nums]
    dups = [n for n in nums if nums.count(n)>1]
    print(f'ch{ch}: count={len(nums)} nums={nums} missing={missing} dups={set(dups) if dups else []}')

print('=== TABLE NUMBER CONTINUITY ===')
by_ch = {}
for t in tabs:
    by_ch.setdefault(t[1], []).append(t[2])
for ch, nums in sorted(by_ch.items()):
    expect = list(range(1, max(nums)+1))
    missing = [n for n in expect if n not in nums]
    dups = [n for n in nums if nums.count(n)>1]
    print(f'ch{ch}: count={len(nums)} nums={nums} missing={missing} dups={set(dups) if dups else []}')

# 2) body references vs captions
body = '\n'.join((p.text or '') for p in doc.paragraphs)
fig_refs = set(re.findall(r'如图(\d+-\d+)所示', body))
tab_refs = set(re.findall(r'如表(\d+-\d+)所示', body))
fig_caps = set(f'{a}-{b}' for _,a,b,_,_ in figs)
tab_caps = set(f'{a}-{b}' for _,a,b,_,_ in tabs)
print('body fig refs not in captions', sorted(fig_refs - fig_caps))
print('captions never referenced', sorted(fig_caps - fig_refs))
print('body tab refs not in captions', sorted(tab_refs - tab_caps))
print('tab captions never referenced', sorted(tab_caps - tab_refs))

# 3) image immediately before each figure caption
from docx.oxml.ns import qn
missing_img = []
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if not re.match(r'^图\d+-\d+', t):
        continue
    ok=False
    for j in (i-1, i-2):
        if j>=0:
            xml = doc.paragraphs[j]._element.xml
            if 'w:drawing' in xml or 'w:pict' in xml:
                ok=True
                break
    if not ok:
        missing_img.append(t)
print('figures without nearby image', missing_img)

# 4) table captions followed by tables (next non-empty is table? hard) count tables
print('docx.tables count', len(doc.tables))
print('table captions', len(tabs))

# 5) compare with reference expected set for finance thesis vs green plant
print('ALL FIGS:')
for f in figs:
    print(' ', f[4])
print('ALL TABS:')
for t in tabs:
    print(' ', t[4])

# 6) sample body ref format remaining spaced forms
spaced = []
for i,p in enumerate(doc.paragraphs):
    t=p.text or ''
    if re.search(r'如图\s+\d|如表\s+\d|图\s+\d+-\d+|表\s+\d+-\d+', t) and not re.match(r'^[图表]\d+-\d+', t.strip()):
        # captions already no space; body leftover
        if re.search(r'图\s+\d+-\d+|表\s+\d+-\d+|如图\s+\d|如表\s+\d', t):
            spaced.append((i,t[:120]))
print('remaining spaced body patterns', len(spaced))
for s in spaced[:20]:
    print(s)
