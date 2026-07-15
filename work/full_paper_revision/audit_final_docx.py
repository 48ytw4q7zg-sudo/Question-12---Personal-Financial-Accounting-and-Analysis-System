from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


FIGURE_EXPECTED = {
    2: list(range(1, 5)),
    4: list(range(1, 22)),
    5: list(range(1, 13)),
}
TABLE_EXPECTED = {
    2: [1],
    3: [1, 2],
    4: list(range(1, 9)),
    6: list(range(1, 7)),
}
CHAPTER_TITLES = (
    "1 绪论",
    "2 可行性分析与需求分析",
    "3 系统相关理论和技术",
    "4 系统设计",
    "5 系统实现",
    "6 系统测试",
    "7 总结和展望",
)


def caption_ids(document: Document, marker: str) -> tuple[list[str], list[str]]:
    ids: list[str] = []
    captions: list[str] = []
    pattern = re.compile(rf"^{marker}(\d+-\d+)\s+\S")
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = pattern.match(text)
        if match:
            ids.append(match.group(1))
            captions.append(text)
    return ids, captions


def expected_ids(spec: dict[int, list[int]]) -> list[str]:
    return [f"{chapter}-{number}" for chapter, numbers in spec.items() for number in numbers]


def body_reference_ids(document: Document, marker: str) -> set[str]:
    references: set[str] = set()
    captions = re.compile(rf"^{marker}\d+-\d+\s")
    direct = re.compile(rf"{marker}(\d+-\d+)")
    range_pattern = re.compile(rf"{marker}(\d+)-(\d+)至{marker}(\d+)-(\d+)")
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or captions.match(text):
            continue
        references.update(direct.findall(text))
        for first_chapter, first_number, last_chapter, last_number in range_pattern.findall(text):
            if first_chapter == last_chapter:
                references.update(
                    f"{first_chapter}-{number}"
                    for number in range(int(first_number), int(last_number) + 1)
                )
    return references


def has_drawing(paragraph) -> bool:
    xml = paragraph._element.xml
    return "w:drawing" in xml or "w:pict" in xml


def has_page_break(paragraph) -> bool:
    if paragraph.paragraph_format.page_break_before is True:
        return True
    return bool(paragraph._element.xpath(".//w:br[@w:type='page']"))


def table_border_value(table, edge: str) -> str | None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        return None
    node = borders.find(qn(f"w:{edge}"))
    return None if node is None else node.get(qn("w:val"))


def audit(path: Path) -> list[str]:
    document = Document(str(path))
    errors: list[str] = []

    figure_ids, figure_captions = caption_ids(document, "图")
    table_ids, table_captions = caption_ids(document, "表")
    expected_figures = expected_ids(FIGURE_EXPECTED)
    expected_tables = expected_ids(TABLE_EXPECTED)

    if figure_ids != expected_figures:
        errors.append(f"figure sequence mismatch: {figure_ids}")
    if table_ids != expected_tables:
        errors.append(f"table sequence mismatch: {table_ids}")
    if len(document.inline_shapes) != len(expected_figures):
        errors.append(
            f"inline shape count {len(document.inline_shapes)} != {len(expected_figures)}"
        )
    if len(document.tables) != len(expected_tables):
        errors.append(f"table count {len(document.tables)} != {len(expected_tables)}")

    figure_refs = body_reference_ids(document, "图")
    table_refs = body_reference_ids(document, "表")
    missing_figure_refs = sorted(set(expected_figures) - figure_refs)
    missing_table_refs = sorted(set(expected_tables) - table_refs)
    if missing_figure_refs:
        errors.append(f"figures without body references: {missing_figure_refs}")
    if missing_table_refs:
        errors.append(f"tables without body references: {missing_table_refs}")

    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if re.match(r"^图\d+-\d+\s", text):
            nearby = paragraphs[max(0, index - 2) : index]
            if not any(has_drawing(item) for item in nearby):
                errors.append(f"figure caption without nearby image: {text}")
        if re.match(r"^[图表]\d+-\d+\s", text):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append(f"caption is not centered: {text}")

    body_children = list(document.element.body.iterchildren())
    for index, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        texts = child.xpath(".//w:t/text()")
        text = "".join(texts).strip()
        if re.match(r"^表\d+-\d+\s", text):
            next_nonempty = None
            for following in body_children[index + 1 :]:
                if following.tag == qn("w:p"):
                    following_text = "".join(following.xpath(".//w:t/text()") ).strip()
                    if not following_text and not following.xpath(".//w:drawing|.//w:pict"):
                        continue
                next_nonempty = following
                break
            if next_nonempty is None or next_nonempty.tag != qn("w:tbl"):
                errors.append(f"table caption not followed by a table: {text}")

    for number, table in enumerate(document.tables, 1):
        for edge in ("left", "right", "insideH", "insideV"):
            if table_border_value(table, edge) not in ("nil", "none"):
                errors.append(f"table {number} has non-three-line border {edge}")
        for edge in ("top", "bottom"):
            if table_border_value(table, edge) != "single":
                errors.append(f"table {number} missing {edge} border")

    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for pattern in (
        r"图\s+\d+-\d+",
        r"表\s+\d+-\d+",
        r"如图\s+\d+-\d+",
        r"如表\s+\d+-\d+",
    ):
        if re.search(pattern, all_text):
            errors.append(f"spaced figure/table format remains: {pattern}")

    required_headings = {
        "1 绪论",
        "1.1 研究背景及意义",
        "1.1.1 研究背景",
        "1.1.2 研究意义",
        "1.2 主要研究内容",
        "2 可行性分析与需求分析",
        "2.1 可行性分析",
        "2.1.1 经济可行性",
        "2.1.2 操作可行性",
        "2.1.3 技术可行性",
        "2.2 系统需求分析",
        "2.2.1 业务需求分析",
        "2.2.2 用户需求分析",
        "2.2.3 性能与非功能需求分析",
        "3 系统相关理论和技术",
        "3.1 开发与运行环境",
        "3.1.1 前端开发环境",
        "3.1.2 后端与数据库环境",
        "3.2 开发框架与关键技术",
        "3.2.1 Vue 3 与前端生态",
        "3.2.2 Spring Boot 与 MyBatis-Plus",
        "3.2.3 MySQL 与事务处理",
        "3.2.4 身份认证与统一异常处理",
        "3.2.5 文件导入、图表与汇率服务",
        "3.3 本章小结",
        "4 系统设计",
        "4.4 系统数据库设计",
        "4.5 详细模块设计",
        "5 系统实现",
        "6 系统测试",
        "7.1 总结",
        "7.2 展望",
    }
    paragraph_texts = {paragraph.text.strip() for paragraph in document.paragraphs}
    missing_headings = sorted(required_headings - paragraph_texts)
    if missing_headings:
        errors.append(f"missing headings: {missing_headings}")

    for title in CHAPTER_TITLES:
        matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == title]
        if len(matches) != 1:
            errors.append(f"chapter heading count mismatch: {title} ({len(matches)})")
            continue
        paragraph = matches[0]
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            errors.append(f"chapter heading is not centered: {title}")
        if paragraph.paragraph_format.page_break_before is not True:
            errors.append(f"chapter heading does not start a new page: {title}")
        for run in paragraph.runs:
            if run.font.size is None or abs(run.font.size.pt - 16) > 0.1:
                errors.append(f"chapter heading font size mismatch: {title}")
                break

    for text in ("表3-2 系统主要技术与版本", "4.5.3 核心模块职责设计"):
        paragraph = next(
            (item for item in document.paragraphs if item.text.strip() == text),
            None,
        )
        if paragraph is None:
            errors.append(f"pagination anchor missing: {text}")
            continue
        previous = paragraph._p.getprevious()
        while previous is not None and previous.tag != qn("w:p"):
            previous = previous.getprevious()
        if previous is None or not has_page_break(type(paragraph)(previous, paragraph._parent)):
            errors.append(f"pagination anchor lacks preceding page break: {text}")

    forbidden_headings = {
        "致谢",
        "参考文献",
        "国内外研究现状",
        "论文组织架构",
        "1.3 国内外研究现状",
        "1.4 论文组织架构",
    }
    present_forbidden = sorted(forbidden_headings & paragraph_texts)
    if present_forbidden:
        errors.append(f"forbidden headings present: {present_forbidden}")

    for placeholder in ("截图可插入位置", "__END_ANCHOR__", "TODO", "待补充"):
        if placeholder in all_text:
            errors.append(f"placeholder remains: {placeholder}")

    required_evidence = (
        "关键词：个人财务；收支记账；预算预警；Spring Boot；Vue 3",
        "budget_alert",
        "BudgetScheduler",
        "OpenCSV",
        "Vue 3",
        "Spring Boot",
        "MyBatis-Plus",
        "175 项测试",
        "174 项通过",
        "2026-08-01",
    )
    for text in required_evidence:
        if text not in all_text:
            errors.append(f"required implementation evidence missing: {text}")

    print(f"file={path}")
    print(f"paragraphs={len(document.paragraphs)}")
    print(f"figures={len(figure_ids)} inline_shapes={len(document.inline_shapes)}")
    print(f"tables={len(table_ids)} table_objects={len(document.tables)}")
    print(f"figure_refs={len(figure_refs)} table_refs={len(table_refs)}")
    print(f"figure_captions={figure_captions}")
    print(f"table_captions={table_captions}")
    if errors:
        print("AUDIT_FAILED")
        for error in errors:
            print(f"- {error}")
    else:
        print("AUDIT_OK")
    return errors


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("document", type=Path)
    args = parser.parse_args()
    raise SystemExit(1 if audit(args.document) else 0)


if __name__ == "__main__":
    main()
